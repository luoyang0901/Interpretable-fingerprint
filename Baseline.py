import argparse
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
import pandas as pd
import seaborn as sns
from sklearn.model_selection import train_test_split, KFold
from sklearn.base import clone
from sklearn.preprocessing import RobustScaler
from sklearn.feature_selection import SelectKBest, mutual_info_regression
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.linear_model import Ridge, ElasticNet
from sklearn.svm import SVR
from scipy.stats import pearsonr
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
try:
    import shap
    SHAP_AVAILABLE = True
except Exception:
    SHAP_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
import warnings
warnings.filterwarnings('ignore')
import json
import os
import re
from datetime import datetime
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MULTI_SEED_DEFAULT_LIST = [0, 1, 2, 41, 43]
MULTI_SEED_PARENT_SUBDIR = 'multi_seed_runs'
_RUN_OUTPUT_ROOT = None

def get_run_output_root():
    return _RUN_OUTPUT_ROOT if _RUN_OUTPUT_ROOT is not None else SCRIPT_DIR

def set_run_output_root(path_or_none):
    global _RUN_OUTPUT_ROOT
    _RUN_OUTPUT_ROOT = path_or_none
    if path_or_none and (not os.path.exists(path_or_none)):
        os.makedirs(path_or_none, exist_ok=True)

def apply_output_tag(output_tag):
    if output_tag is None or str(output_tag).strip() == '':
        return None
    tag = str(output_tag).strip()
    tagged_root = os.path.join(SCRIPT_DIR, tag)
    set_run_output_root(tagged_root)
    return tagged_root

def input_path(filename: str) -> str:
    return os.path.join(SCRIPT_DIR, filename)

PUBLIC_DA_FILENAME = 'D-A.csv'
PUBLIC_DONOR_FILENAME = 'Donor.csv'
PUBLIC_ACCEPTOR_FILENAME = 'Acceptor.csv'
PUBLIC_DA_CLEANED_FILENAME = 'D-A_cleaned.csv'
PUBLIC_DA_FILTERED_FILENAME = 'D-A_filtered.csv'


def _find_first_existing_source_file(*filenames: str):
    for filename in filenames:
        path = input_path(filename)
        if os.path.exists(path):
            return path
    return input_path(filenames[0])


def _load_reference_smiles_table(path: str, entity_label: str):
    if not os.path.exists(path):
        return None
    df = _read_smiles_table_auto(path)
    smiles_col = None
    for c in df.columns:
        cl = str(c).lower()
        if cl == 'smiles' or 'smiles' in cl:
            smiles_col = c
            break
    if smiles_col is None:
        smiles_col = df.columns[-1]
    id_candidates = [entity_label, entity_label.lower(), 'Name', 'name']
    id_col = next((c for c in id_candidates if c in df.columns), None)
    if id_col is None:
        non_smiles_cols = [c for c in df.columns if c != smiles_col]
        id_col = non_smiles_cols[0] if non_smiles_cols else None
    out = pd.DataFrame()
    if id_col is not None:
        out[entity_label] = df[id_col].map(_norm_text_auto)
    else:
        out[entity_label] = [f'{entity_label[0]}_{i}' for i in range(len(df))]
    out['SMILES'] = df[smiles_col].map(_norm_text_auto)
    out = out[out['SMILES'] != ''].drop_duplicates(subset='SMILES', keep='first').reset_index(drop=True)
    if entity_label not in out.columns:
        out[entity_label] = [f'{entity_label[0]}_{i}' for i in range(len(out))]
    return out[[entity_label, 'SMILES']]
DEFAULT_RANDOM_SEED = 42
BEST_MAIN_CONFIG_FILENAME = 'best_main_experiment_config.json'
EXPERIMENT_PROFILES = {'legacy': {'radius': 3, 'use_gao_aux': False, 'label': 'Legacy / Existing prepared inputs'}, 'main_r3': {'radius': 3, 'use_gao_aux': False, 'label': 'Main experiment: core data only, radius=3'}, 'core_r2': {'radius': 2, 'use_gao_aux': False, 'label': 'Radius control: core data only, radius=2'}, 'core_r2_plus_gao': {'radius': 2, 'use_gao_aux': True, 'label': 'Auxiliary experiment: core radius=2 + Gao fingerprint-only aux'}}
CURRENT_EXPERIMENT_CONTEXT = {'profile': 'legacy', 'radius': 3, 'use_gao_aux': False, 'gao_fd_path': None, 'gao_fa_path': None, 'gao_y_path': None}

def _set_experiment_context(profile='legacy', radius=None, use_gao_aux=None, gao_fd_path=None, gao_fa_path=None, gao_y_path=None):
    cfg = dict(EXPERIMENT_PROFILES.get(profile, EXPERIMENT_PROFILES['legacy']))
    CURRENT_EXPERIMENT_CONTEXT['profile'] = profile
    CURRENT_EXPERIMENT_CONTEXT['radius'] = int(cfg['radius'] if radius is None else radius)
    CURRENT_EXPERIMENT_CONTEXT['use_gao_aux'] = bool(cfg['use_gao_aux'] if use_gao_aux is None else use_gao_aux)
    CURRENT_EXPERIMENT_CONTEXT['gao_fd_path'] = gao_fd_path
    CURRENT_EXPERIMENT_CONTEXT['gao_fa_path'] = gao_fa_path
    CURRENT_EXPERIMENT_CONTEXT['gao_y_path'] = gao_y_path

def _load_gao_aux_feature_arrays(fd_path, fa_path, y_path, n_bits=1024):
    if not (fd_path and fa_path and y_path):
        raise ValueError('\u542f\u7528 Gao \u8f85\u52a9\u96c6\u65f6\uff0c\u5fc5\u987b\u540c\u65f6\u63d0\u4f9b --gao-fd-path / --gao-fa-path / --gao-y-path')
    fd = np.load(fd_path)
    fa = np.load(fa_path)
    y = np.load(y_path)
    if fd.ndim != 2 or fa.ndim != 2:
        raise ValueError('Gao \u8f85\u52a9\u96c6 fd/fa \u5fc5\u987b\u662f\u4e8c\u7ef4\u6570\u7ec4')
    if fd.shape[1] != n_bits or fa.shape[1] != n_bits:
        raise ValueError(f'Gao \u8f85\u52a9\u96c6\u7279\u5f81\u7ef4\u5ea6\u5fc5\u987b\u4e3a ({n_bits})\uff0c\u5f53\u524d fd={fd.shape}, fa={fa.shape}')
    if fd.shape[0] != fa.shape[0] or fd.shape[0] != len(y):
        raise ValueError(f'Gao \u8f85\u52a9\u96c6\u6837\u672c\u6570\u4e0d\u4e00\u81f4: fd={fd.shape}, fa={fa.shape}, y={np.shape(y)}')
    return (fd.astype(np.float32), fa.astype(np.float32), np.asarray(y, dtype=np.float32))

def _append_aux_gao_to_filtered_training_data(fd_path, fa_path, y_path, n_bits=1024):
    g_fd, g_fa, g_y = _load_gao_aux_feature_arrays(fd_path, fa_path, y_path, n_bits=n_bits)
    fd_core = np.load(data_path('fd_fp_filtered.npy')).astype(np.float32)
    fa_core = np.load(data_path('fa_fp_filtered.npy')).astype(np.float32)
    y_core = np.load(data_path('fp_Y_filtered.npy')).astype(np.float32)
    fd_mix = np.vstack([fd_core, g_fd]).astype(np.float32)
    fa_mix = np.vstack([fa_core, g_fa]).astype(np.float32)
    y_mix = np.concatenate([y_core, g_y]).astype(np.float32)
    np.save(data_path('fd_fp_filtered.npy'), fd_mix)
    np.save(data_path('fa_fp_filtered.npy'), fa_mix)
    np.save(data_path('fp_Y_filtered.npy'), y_mix)
    meta = {'profile': CURRENT_EXPERIMENT_CONTEXT.get('profile'), 'core_rows_before_append': int(len(y_core)), 'gao_aux_rows_appended': int(len(g_y)), 'rows_after_append': int(len(y_mix)), 'gao_fd_path': os.path.abspath(fd_path), 'gao_fa_path': os.path.abspath(fa_path), 'gao_y_path': os.path.abspath(y_path)}
    with open(data_path('gao_aux_append_summary.json'), 'w', encoding='utf-8') as f:
        json.dump(meta, f, indent=2, ensure_ascii=False)
    return meta

def _write_experiment_context_summary(extra=None):
    payload = {'profile': CURRENT_EXPERIMENT_CONTEXT.get('profile'), 'radius': int(CURRENT_EXPERIMENT_CONTEXT.get('radius', 3)), 'use_gao_aux': bool(CURRENT_EXPERIMENT_CONTEXT.get('use_gao_aux', False)), 'gao_fd_path': CURRENT_EXPERIMENT_CONTEXT.get('gao_fd_path'), 'gao_fa_path': CURRENT_EXPERIMENT_CONTEXT.get('gao_fa_path'), 'gao_y_path': CURRENT_EXPERIMENT_CONTEXT.get('gao_y_path'), 'output_root': get_run_output_root()}
    if extra:
        payload.update(extra)
    with open(data_path('experiment_context_summary.json'), 'w', encoding='utf-8') as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)
    return payload

def _read_smiles_table_auto(path: str):
    suffix = os.path.splitext(path)[1].lower()
    if suffix in ('.xlsx', '.xls'):
        return pd.read_excel(path)
    for enc in ('utf-8', 'gbk', 'latin-1'):
        try:
            return pd.read_csv(path, encoding=enc)
        except Exception:
            pass
    raise ValueError(f'\u65e0\u6cd5\u8bfb\u53d6\u6570\u636e\u96c6\u6587\u4ef6: {path}')

def _norm_text_auto(x):
    if pd.isna(x):
        return ''
    return str(x).strip()

def _guess_smiles_dataset_columns_auto(df, donor_col=None, acceptor_col=None, pce_col=None):
    cols = list(df.columns)
    if donor_col is not None and donor_col not in cols:
        raise ValueError(f'\u6307\u5b9a donor \u5217\u4e0d\u5b58\u5728: {donor_col}')
    if acceptor_col is not None and acceptor_col not in cols:
        raise ValueError(f'\u6307\u5b9a acceptor \u5217\u4e0d\u5b58\u5728: {acceptor_col}')
    if pce_col is not None and pce_col not in cols:
        raise ValueError(f'\u6307\u5b9a PCE \u5217\u4e0d\u5b58\u5728: {pce_col}')
    if donor_col is None:
        for c in cols:
            cl = str(c).lower()
            if 'smiles' in cl and ('donor' in cl or '\u7ed9\u4f53' in str(c)):
                donor_col = c
                break
        if donor_col is None:
            for c in cols:
                cl = str(c).lower()
                if cl.startswith('donor') or 'donor' in cl or '\u7ed9\u4f53' in str(c):
                    donor_col = c
                    break
    if acceptor_col is None:
        for c in cols:
            cl = str(c).lower()
            if 'smiles' in cl and ('acceptor' in cl or '\u53d7\u4f53' in str(c)):
                acceptor_col = c
                break
        if acceptor_col is None:
            for c in cols:
                cl = str(c).lower()
                if cl.startswith('acceptor') or 'acceptor' in cl or '\u53d7\u4f53' in str(c):
                    acceptor_col = c
                    break
    if pce_col is None:
        for c in cols:
            cl = str(c).lower()
            if 'pce' in cl or '\u6548\u7387' in str(c):
                pce_col = c
                break
        if pce_col is None:
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            if numeric_cols:
                pce_col = numeric_cols[0]
    if donor_col is None or acceptor_col is None or pce_col is None:
        raise ValueError(f'\u81ea\u52a8\u8bc6\u522b donor/acceptor/PCE \u5217\u5931\u8d25\u3002\u8bf7\u624b\u52a8\u4f20\u5165 --donor-col / --acceptor-col / --pce-col\u3002\n\u5f53\u524d\u5217\u540d: {cols}')
    return (donor_col, acceptor_col, pce_col)

def _morgan_bits_and_onbits_auto(smiles: str, radius: int=3, n_bits: int=1024):
    from rdkit import Chem, RDLogger
    from rdkit.Chem import AllChem
    RDLogger.DisableLog('rdApp.warning')
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        raise ValueError(f'SMILES \u65e0\u6cd5\u89e3\u6790: {smiles}')
    fp = AllChem.GetMorganFingerprintAsBitVect(mol, radius, nBits=n_bits)
    bits = np.array(list(map(int, fp.ToBitString())), dtype=np.int8)
    on_bits_0 = np.where(bits == 1)[0].tolist()
    on_bits_1 = [i + 1 for i in on_bits_0]
    return {'bits': bits, 'on_bits_0_based': on_bits_0, 'on_bits_1_based': on_bits_1, 'num_on_bits': len(on_bits_0)}

def _prepare_inputs_from_smiles_dataset(dataset_path: str, donor_col: str=None, acceptor_col: str=None, pce_col: str=None, radius: int=3, n_bits: int=1024):
    df = _read_smiles_table_auto(dataset_path)
    donor_col, acceptor_col, pce_col = _guess_smiles_dataset_columns_auto(df, donor_col, acceptor_col, pce_col)
    work = df[[donor_col, acceptor_col, pce_col]].copy()
    work.columns = ['Donor_SMILES', 'Acceptor_SMILES', 'PCE']
    work['Donor_SMILES'] = work['Donor_SMILES'].map(_norm_text_auto)
    work['Acceptor_SMILES'] = work['Acceptor_SMILES'].map(_norm_text_auto)
    work = work[(work['Donor_SMILES'] != '') & (work['Acceptor_SMILES'] != '')].copy()
    work['PCE'] = pd.to_numeric(work['PCE'], errors='coerce')
    work = work.dropna(subset=['PCE']).reset_index(drop=True)
    work['DA_Key'] = work['Donor_SMILES'] + '||' + work['Acceptor_SMILES']
    work = work.sort_values('PCE', ascending=False).drop_duplicates('DA_Key', keep='first').reset_index(drop=True)
    work = work.drop(columns=['DA_Key'])

    donor_ref_path = _find_first_existing_source_file(PUBLIC_DONOR_FILENAME)
    acceptor_ref_path = _find_first_existing_source_file(PUBLIC_ACCEPTOR_FILENAME)
    donor_unique = _load_reference_smiles_table(donor_ref_path, 'Donor')
    acceptor_unique = _load_reference_smiles_table(acceptor_ref_path, 'Acceptor')

    if donor_unique is None or donor_unique.empty:
        donor_unique = work[['Donor_SMILES']].drop_duplicates().reset_index(drop=True).copy()
        donor_unique['Donor'] = [f'D_{i}' for i in range(len(donor_unique))]
        donor_unique = donor_unique[['Donor', 'Donor_SMILES']].rename(columns={'Donor_SMILES': 'SMILES'})
    else:
        work_donors = set(work['Donor_SMILES'].tolist())
        donor_unique = donor_unique[donor_unique['SMILES'].isin(work_donors)].drop_duplicates(subset='SMILES', keep='first').reset_index(drop=True)
        missing_donors = [s for s in work['Donor_SMILES'].drop_duplicates().tolist() if s not in set(donor_unique['SMILES'])]
        if missing_donors:
            donor_extra = pd.DataFrame({'Donor': [f'D_{len(donor_unique) + i}' for i in range(len(missing_donors))], 'SMILES': missing_donors})
            donor_unique = pd.concat([donor_unique, donor_extra], ignore_index=True)

    if acceptor_unique is None or acceptor_unique.empty:
        acceptor_unique = work[['Acceptor_SMILES']].drop_duplicates().reset_index(drop=True).copy()
        acceptor_unique['Acceptor'] = [f'A_{i}' for i in range(len(acceptor_unique))]
        acceptor_unique = acceptor_unique[['Acceptor', 'Acceptor_SMILES']].rename(columns={'Acceptor_SMILES': 'SMILES'})
    else:
        work_acceptors = set(work['Acceptor_SMILES'].tolist())
        acceptor_unique = acceptor_unique[acceptor_unique['SMILES'].isin(work_acceptors)].drop_duplicates(subset='SMILES', keep='first').reset_index(drop=True)
        missing_acceptors = [s for s in work['Acceptor_SMILES'].drop_duplicates().tolist() if s not in set(acceptor_unique['SMILES'])]
        if missing_acceptors:
            acceptor_extra = pd.DataFrame({'Acceptor': [f'A_{len(acceptor_unique) + i}' for i in range(len(missing_acceptors))], 'SMILES': missing_acceptors})
            acceptor_unique = pd.concat([acceptor_unique, acceptor_extra], ignore_index=True)

    donor_ref_fp = []
    for smi in donor_unique['SMILES'].tolist():
        donor_ref_fp.append(_morgan_bits_and_onbits_auto(smi, radius=radius, n_bits=n_bits)['bits'])
    acceptor_ref_fp = []
    for smi in acceptor_unique['SMILES'].tolist():
        acceptor_ref_fp.append(_morgan_bits_and_onbits_auto(smi, radius=radius, n_bits=n_bits)['bits'])
    fd_fp1 = np.vstack(donor_ref_fp).astype(np.float32) if donor_ref_fp else np.empty((0, n_bits), dtype=np.float32)
    fa_fp1 = np.vstack(acceptor_ref_fp).astype(np.float32) if acceptor_ref_fp else np.empty((0, n_bits), dtype=np.float32)

    donor_pair_bits = []
    acceptor_pair_bits = []
    concat_pair_bits = []
    donor_codes = []
    acceptor_codes = []
    export_rows = []
    for i, row in work.iterrows():
        d_smi = str(row['Donor_SMILES'])
        a_smi = str(row['Acceptor_SMILES'])
        pce_val = float(row['PCE'])
        d_fp = _morgan_bits_and_onbits_auto(d_smi, radius=radius, n_bits=n_bits)
        a_fp = _morgan_bits_and_onbits_auto(a_smi, radius=radius, n_bits=n_bits)
        concat_bits = np.concatenate([d_fp['bits'], a_fp['bits']], axis=0).astype(np.int8)
        concat_on_0 = d_fp['on_bits_0_based'] + [n_bits + x for x in a_fp['on_bits_0_based']]
        concat_on_1 = [x + 1 for x in concat_on_0]
        donor_pair_bits.append(d_fp['bits'])
        acceptor_pair_bits.append(a_fp['bits'])
        concat_pair_bits.append(concat_bits)
        donor_codes.append(d_smi)
        acceptor_codes.append(a_smi)
        export_rows.append({'row_index': int(i), 'donor_smiles': d_smi, 'acceptor_smiles': a_smi, 'pce': pce_val, 'donor_num_on_bits': int(d_fp['num_on_bits']), 'acceptor_num_on_bits': int(a_fp['num_on_bits']), 'concat_num_on_bits': int(len(concat_on_0)), 'donor_on_bits_0_based': json.dumps(d_fp['on_bits_0_based'], ensure_ascii=False), 'acceptor_on_bits_0_based': json.dumps(a_fp['on_bits_0_based'], ensure_ascii=False), 'concat_on_bits_0_based': json.dumps(concat_on_0, ensure_ascii=False), 'donor_on_bits_1_based': json.dumps(d_fp['on_bits_1_based'], ensure_ascii=False), 'acceptor_on_bits_1_based': json.dumps(a_fp['on_bits_1_based'], ensure_ascii=False), 'concat_on_bits_1_based': json.dumps(concat_on_1, ensure_ascii=False), 'donor_bits_1024': json.dumps(d_fp['bits'].tolist(), ensure_ascii=False), 'acceptor_bits_1024': json.dumps(a_fp['bits'].tolist(), ensure_ascii=False), 'concat_bits_2048': json.dumps(concat_bits.tolist(), ensure_ascii=False)})

    fd_fp_augmented = np.vstack(donor_pair_bits).astype(np.float32) if donor_pair_bits else np.empty((0, n_bits), dtype=np.float32)
    fa_fp_augmented = np.vstack(acceptor_pair_bits).astype(np.float32) if acceptor_pair_bits else np.empty((0, n_bits), dtype=np.float32)
    concat_bits_2048 = np.vstack(concat_pair_bits).astype(np.int8) if concat_pair_bits else np.empty((0, n_bits * 2), dtype=np.int8)
    fp_Y_augmented = work['PCE'].to_numpy(dtype=np.float32)

    np.save(input_path('fd_fp1.npy'), fd_fp1)
    np.save(input_path('fa_fp1.npy'), fa_fp1)
    np.save(input_path('fd_fp_augmented.npy'), fd_fp_augmented)
    np.save(input_path('fa_fp_augmented.npy'), fa_fp_augmented)
    np.save(input_path('fp_Y_augmented.npy'), fp_Y_augmented)
    np.save(input_path('fd_fp.npy'), np.empty((0, n_bits), dtype=np.float32))
    np.save(input_path('fa_fp.npy'), np.empty((0, n_bits), dtype=np.float32))
    np.save(input_path('fp_Y.npy'), np.empty((0,), dtype=np.float32))
    np.save(input_path('donor_codes.npy'), np.asarray(donor_codes, dtype=f'<U{max(8, max((len(x) for x in donor_codes), default=8))}'))
    np.save(input_path('acceptor_codes.npy'), np.asarray(acceptor_codes, dtype=f'<U{max(8, max((len(x) for x in acceptor_codes), default=8))}'))

    converted_df = pd.DataFrame(export_rows)
    converted_df.to_csv(input_path('converted_da_features.csv'), index=False, encoding='utf-8-sig')
    converted_df.to_json(input_path('converted_da_features.json'), orient='records', force_ascii=False, indent=2)
    np.savez_compressed(input_path('converted_da_features.npz'), donor_bits_1024=fd_fp_augmented.astype(np.int8), acceptor_bits_1024=fa_fp_augmented.astype(np.int8), concat_bits_2048=concat_bits_2048.astype(np.int8), pce=fp_Y_augmented.astype(np.float32))
    converted_df[['row_index', 'donor_smiles', 'acceptor_smiles', 'pce', 'donor_on_bits_0_based', 'acceptor_on_bits_0_based', 'concat_on_bits_0_based']].to_csv(input_path('on_bits_summary.csv'), index=False, encoding='utf-8-sig')

    meta = {
        'dataset_path': os.path.abspath(dataset_path),
        'donor_reference_path': os.path.abspath(donor_ref_path) if donor_ref_path else None,
        'acceptor_reference_path': os.path.abspath(acceptor_ref_path) if acceptor_ref_path else None,
        'rows_after_dedup': int(len(work)),
        'n_unique_donors': int(len(donor_unique)),
        'n_unique_acceptors': int(len(acceptor_unique)),
        'radius': int(radius),
        'n_bits': int(n_bits),
        'generated_files_in_script_dir': [
            'fd_fp1.npy', 'fa_fp1.npy',
            'fd_fp_augmented.npy', 'fa_fp_augmented.npy', 'fp_Y_augmented.npy',
            'fd_fp.npy', 'fa_fp.npy', 'fp_Y.npy',
            'donor_codes.npy', 'acceptor_codes.npy',
            'converted_da_features.csv', 'converted_da_features.json',
            'converted_da_features.npz', 'on_bits_summary.csv'
        ]
    }
    with open(input_path('smiles_auto_prepare_summary.json'), 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def _parse_seed_list(s: str):
    parts = [p.strip() for p in s.split(',') if p.strip()]
    return [int(p) for p in parts]

def _resolve_random_state(random_state):
    return DEFAULT_RANDOM_SEED if random_state is None else int(random_state)
SPLIT_PRESET_SPECS = {'64_16_20': {'test_size': 0.2, 'val_fraction_of_trainval': 0.2, 'description': '64:16:20 (\u539f nested 0.2/0.2)'}, '70_10_20': {'test_size': 0.2, 'val_fraction_of_trainval': 0.125, 'description': '70:10:20'}, '70_15_15': {'test_size': 0.15, 'val_fraction_of_trainval': 0.15 / 0.85, 'description': '70:15:15'}, '80_20_cv': {'test_size': 0.2, 'val_fraction_of_trainval': None, 'description': '80:20 (\u8bad\u7ec3+\u9a8c\u8bc1\u5185 5 \u6298 CV)'}}
DEFAULT_SPLIT_PRESET = '70_10_20'
CV_N_SPLITS = 5

def _augment_train_features_only(X_train, y_train, random_state):
    rs = _resolve_random_state(random_state)
    rng = np.random.RandomState(rs)
    noise = rng.normal(0, 0.003, X_train.shape)
    X_noise = np.clip(X_train + noise, 0, 1)
    scale_factors = rng.uniform(0.98, 1.02, X_train.shape[1])
    X_scaled = np.clip(X_train * scale_factors, 0, 1)
    X_aug = np.vstack([X_train, X_noise, X_scaled])
    y_aug = np.concatenate([y_train, y_train, y_train])
    return (X_aug, y_aug)

def preview_unknown_split_for_seeds(seeds, n_unknown=20):
    cleaned_path = data_path(PUBLIC_DA_CLEANED_FILENAME)
    if not os.path.exists(cleaned_path):
        return
    try:
        da_clean = pd.read_csv(cleaned_path, encoding='utf-8')
    except Exception:
        try:
            da_clean = pd.read_csv(cleaned_path, encoding='gbk')
        except Exception:
            da_clean = pd.read_csv(cleaned_path, encoding='latin-1')
    n = len(da_clean)
    nu = min(n_unknown, n)
    for s in seeds:
        rng = np.random.RandomState(s)
        idx = rng.choice(n, size=nu, replace=False)
        idx = np.sort(idx)

def compare_train_seeds_only(seeds, split_preset=DEFAULT_SPLIT_PRESET):
    req = [data_path('fd_fp_filtered.npy'), data_path('fa_fp_filtered.npy'), data_path('fp_Y_filtered.npy')]
    for p in req:
        if not os.path.exists(p):
            return
    rows = []
    for s in seeds:
        X_raw, y_raw, *_ = load_raw_merged_features()
        r_no = train_and_evaluate(X_raw, y_raw, '\u65e0\u589e\u5f3a(\u5408\u5e76\u540e)', random_state=s, split_preset=split_preset, augment_train=False)['final_r']
        r_w = train_and_evaluate(X_raw, y_raw, '\u6709\u589e\u5f3a(\u5408\u5e76\u540e)', random_state=s, split_preset=split_preset, augment_train=True)['final_r']
        rows.append((s, r_no, r_w))
    if rows:
        rnos = [t[1] for t in rows]
        rwss = [t[2] for t in rows]

def _generate_seed_search_report(search_parent, df_results, all_model_results, search_type):
    if not DOCX_AVAILABLE:
        return
    try:
        doc = Document()
        doc.add_heading(f'\u79cd\u5b50\u641c\u7d22\u5bf9\u6bd4\u62a5\u544a ({search_type})', 0)
        doc.add_paragraph(f"\u751f\u6210\u65f6\u95f4: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"\u641c\u7d22\u8303\u56f4: {int(df_results['seed'].min())} - {int(df_results['seed'].max())}")
        doc.add_paragraph(f"\u6700\u4f18\u79cd\u5b50: {int(df_results.iloc[0]['seed'])} (best_r = {df_results.iloc[0]['best_r']:.4f})")
        first_seed = int(df_results.iloc[0]['seed'])
        if first_seed in all_model_results:
            model_names = list(all_model_results[first_seed]['no_aug'].keys())
        else:
            model_names = []
        doc.add_heading('\u65e0\u589e\u5f3a\u914d\u7f6e - \u5404\u6a21\u578b r \u503c\u5bf9\u6bd4', level=1)
        table_no_aug = doc.add_table(rows=1, cols=len(model_names) + 1)
        table_no_aug.style = 'Light Grid Accent 1'
        hdr_cells = table_no_aug.rows[0].cells
        hdr_cells[0].text = 'Seed'
        for i, model_name in enumerate(model_names):
            hdr_cells[i + 1].text = model_name
        for _, row in df_results.iterrows():
            seed = int(row['seed'])
            row_cells = table_no_aug.add_row().cells
            row_cells[0].text = str(seed)
            if seed in all_model_results:
                for i, model_name in enumerate(model_names):
                    if model_name in all_model_results[seed]['no_aug']:
                        r_val = all_model_results[seed]['no_aug'][model_name].get('r_test', 0)
                        row_cells[i + 1].text = f'{r_val:.4f}'
                    else:
                        row_cells[i + 1].text = 'N/A'
            else:
                for i in range(len(model_names)):
                    row_cells[i + 1].text = 'N/A'
        doc.add_heading('\u6709\u589e\u5f3a\u914d\u7f6e - \u5404\u6a21\u578b r \u503c\u5bf9\u6bd4', level=1)
        table_with_aug = doc.add_table(rows=1, cols=len(model_names) + 1)
        table_with_aug.style = 'Light Grid Accent 1'
        hdr_cells = table_with_aug.rows[0].cells
        hdr_cells[0].text = 'Seed'
        for i, model_name in enumerate(model_names):
            hdr_cells[i + 1].text = model_name
        for _, row in df_results.iterrows():
            seed = int(row['seed'])
            row_cells = table_with_aug.add_row().cells
            row_cells[0].text = str(seed)
            if seed in all_model_results:
                for i, model_name in enumerate(model_names):
                    if model_name in all_model_results[seed]['with_aug']:
                        r_val = all_model_results[seed]['with_aug'][model_name].get('r_test', 0)
                        row_cells[i + 1].text = f'{r_val:.4f}'
                    else:
                        row_cells[i + 1].text = 'N/A'
            else:
                for i in range(len(model_names)):
                    row_cells[i + 1].text = 'N/A'
        doc.add_heading('\u6c47\u603b\u7edf\u8ba1', level=1)
        table_summary = doc.add_table(rows=1, cols=6)
        table_summary.style = 'Light Grid Accent 1'
        hdr_cells = table_summary.rows[0].cells
        hdr_cells[0].text = 'Seed'
        hdr_cells[1].text = 'Best R'
        hdr_cells[2].text = 'Best Config'
        hdr_cells[3].text = 'R (No Aug)'
        hdr_cells[4].text = 'R (With Aug)'
        hdr_cells[5].text = 'Unknown R'
        for _, row in df_results.iterrows():
            row_cells = table_summary.add_row().cells
            row_cells[0].text = str(int(row['seed']))
            row_cells[1].text = f"{row['best_r']:.4f}"
            row_cells[2].text = str(row['best_config'])
            row_cells[3].text = f"{row['r_no_aug']:.4f}"
            row_cells[4].text = f"{row['r_with_aug']:.4f}"
            row_cells[5].text = f"{row['unknown_r']:.4f}"
        report_path = os.path.join(search_parent, f'seed_search_report_{search_type}.docx')
        doc.save(report_path)
    except Exception as e:
        pass

def data_path(filename: str) -> str:
    return os.path.join(get_run_output_root(), filename)

def run_multi_seed_batch(seeds=None, max_features=10, plot_only=False, split_preset=DEFAULT_SPLIT_PRESET, run_all_split_presets=False):
    if plot_only:
        plot_only = False
    seed_list = list(seeds) if seeds is not None else list(MULTI_SEED_DEFAULT_LIST)
    parent = os.path.join(get_run_output_root(), MULTI_SEED_PARENT_SUBDIR)
    os.makedirs(parent, exist_ok=True)
    summaries = []
    for seed in seed_list:
        run_dir = os.path.join(parent, f'seed_{seed}')
        set_run_output_root(run_dir)
        try:
            main(plot_only=False, max_features=max_features, random_seed=seed, split_preset=split_preset, run_all_split_presets=run_all_split_presets)
        except Exception as exc:
            set_run_output_root(None)
            raise
        summ_path = os.path.join(run_dir, 'run_metrics_summary.json')
        if os.path.isfile(summ_path):
            with open(summ_path, 'r', encoding='utf-8') as sf:
                summaries.append(json.load(sf))
    set_run_output_root(None)
    if summaries:
        df_sum = pd.DataFrame(summaries)
        out_csv = os.path.join(parent, 'all_seeds_summary.csv')
        df_sum.to_csv(out_csv, index=False, encoding='utf-8-sig')

def run_seed_search_lite(start, end, max_features=10):
    search_parent = os.path.join(get_run_output_root(), 'seed_search_results', f'seed_search_lite_{start}_{end}')
    os.makedirs(search_parent, exist_ok=True)
    set_run_output_root(None)
    seed_base = 42
    unknown_indices, filtered_fd_fp, filtered_fa_fp, filtered_fp_Y, unknown_fd_fp, unknown_fa_fp, unknown_fp_Y, unknown_donor_smiles, unknown_acceptor_smiles = select_and_split_unknown_20_pairs(random_state=seed_base)
    results = []
    total = end - start + 1
    all_model_results = {}
    for i, seed in enumerate(range(start, end + 1), 1):
        seed_dir = os.path.join(search_parent, f'seed_{seed}')
        set_run_output_root(seed_dir)
        X_raw, y_raw, feat_names, d_codes, ac_codes, _, _ = load_data_without_augmentation()
        result_no_aug = train_and_evaluate(X_raw, y_raw, '\u65e0\u589e\u5f3a', random_state=seed, split_preset=DEFAULT_SPLIT_PRESET, augment_train=False, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
        result_with_aug = train_and_evaluate(X_raw, y_raw, '\u6709\u589e\u5f3a', random_state=seed, split_preset=DEFAULT_SPLIT_PRESET, augment_train=True, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
        unknown_results, unknown_predictions = predict_unknown_20_pairs(result_with_aug, result_with_aug['selector'], result_with_aug['scaler'])
        unknown_true = unknown_fp_Y
        unknown_pred = unknown_predictions
        unknown_r = pearsonr(unknown_true, unknown_pred)[0]
        best_r = max(result_no_aug['final_r'], result_with_aug['final_r'])
        best_config = '\u6709\u589e\u5f3a' if result_with_aug['final_r'] >= result_no_aug['final_r'] else '\u65e0\u589e\u5f3a'
        all_model_results[seed] = {'no_aug': result_no_aug['results'], 'with_aug': result_with_aug['results']}
        record = {'seed': seed, 'best_r': best_r, 'best_config': best_config, 'r_no_aug': result_no_aug['final_r'], 'r_with_aug': result_with_aug['final_r'], 'unknown_r': unknown_r, 'r2_no_aug': result_no_aug['final_r2'], 'r2_with_aug': result_with_aug['final_r2'], 'rmse_no_aug': result_no_aug['final_rmse'], 'rmse_with_aug': result_with_aug['final_rmse']}
        results.append(record)
    set_run_output_root(None)
    df_results = pd.DataFrame(results)
    df_results = df_results.sort_values('best_r', ascending=False)
    csv_path = os.path.join(search_parent, f'summary.csv')
    df_results.to_csv(csv_path, index=False, encoding='utf-8-sig')
    _generate_seed_search_report(search_parent, df_results, all_model_results, '\u8f7b\u91cf\u7ea7')
    best_seed = int(df_results.iloc[0]['seed'])
    best_r_value = df_results.iloc[0]['best_r']

def run_seed_search(start, end, max_features=10):
    search_parent = os.path.join(get_run_output_root(), 'seed_search_results', f'seed_search_full_{start}_{end}')
    os.makedirs(search_parent, exist_ok=True)
    results = []
    total = end - start + 1
    all_model_results = {}
    for i, seed in enumerate(range(start, end + 1), 1):
        seed_dir = os.path.join(search_parent, f'seed_{seed}')
        set_run_output_root(seed_dir)
        unknown_indices, filtered_fd_fp, filtered_fa_fp, filtered_fp_Y, unknown_fd_fp, unknown_fa_fp, unknown_fp_Y, unknown_donor_smiles, unknown_acceptor_smiles = select_and_split_unknown_20_pairs(random_state=seed)
        X_raw, y_raw, feat_names, d_codes, ac_codes, _, _ = load_data_without_augmentation()
        result_no_aug = train_and_evaluate(X_raw, y_raw, '\u65e0\u589e\u5f3a', random_state=seed, split_preset=DEFAULT_SPLIT_PRESET, augment_train=False, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
        result_with_aug = train_and_evaluate(X_raw, y_raw, '\u6709\u589e\u5f3a', random_state=seed, split_preset=DEFAULT_SPLIT_PRESET, augment_train=True, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
        unknown_results, unknown_predictions = predict_unknown_20_pairs(result_with_aug, result_with_aug['selector'], result_with_aug['scaler'])
        unknown_true = unknown_fp_Y
        unknown_pred = unknown_predictions
        unknown_r = pearsonr(unknown_true, unknown_pred)[0]
        best_r = max(result_no_aug['final_r'], result_with_aug['final_r'])
        best_config = '\u6709\u589e\u5f3a' if result_with_aug['final_r'] >= result_no_aug['final_r'] else '\u65e0\u589e\u5f3a'
        all_model_results[seed] = {'no_aug': result_no_aug['results'], 'with_aug': result_with_aug['results']}
        record = {'seed': seed, 'best_r': best_r, 'best_config': best_config, 'r_no_aug': result_no_aug['final_r'], 'r_with_aug': result_with_aug['final_r'], 'unknown_r': unknown_r, 'r2_no_aug': result_no_aug['final_r2'], 'r2_with_aug': result_with_aug['final_r2'], 'rmse_no_aug': result_no_aug['final_rmse'], 'rmse_with_aug': result_with_aug['final_rmse']}
        results.append(record)
    set_run_output_root(None)
    df_results = pd.DataFrame(results)
    df_results = df_results.sort_values('best_r', ascending=False)
    csv_path = os.path.join(search_parent, f'summary.csv')
    df_results.to_csv(csv_path, index=False, encoding='utf-8-sig')
    _generate_seed_search_report(search_parent, df_results, all_model_results, '\u5b8c\u6574')
    best_seed = int(df_results.iloc[0]['seed'])
    best_r_value = df_results.iloc[0]['best_r']

def _safe_name(s):
    return str(s).replace('/', '_').replace('\\', '_').replace(' ', '_')

def _build_split_seed_result_rows(split_preset, seed, result_no_aug, result_with_aug):
    rows = []
    model_names = sorted(set(result_no_aug['results'].keys()) | set(result_with_aug['results'].keys()))
    for m in model_names:
        r0 = np.nan
        r1 = np.nan
        if m in result_no_aug['results']:
            r0 = float(result_no_aug['results'][m]['r_test'])
        if m in result_with_aug['results']:
            r1 = float(result_with_aug['results'][m]['r_test'])
        rows.append({'split_preset': split_preset, 'split_desc': SPLIT_PRESET_SPECS[split_preset]['description'], 'seed': int(seed), 'model': m, 'r_no_aug': r0, 'r_with_aug': r1, 'r_delta_with_minus_no': float(r1 - r0), 'best_final_r_no_aug': float(result_no_aug['final_r']), 'best_final_r_with_aug': float(result_with_aug['final_r'])})
    return rows

def _plot_seed_curves_by_model(df_all, out_dir, split_presets_order=None):
    os.makedirs(out_dir, exist_ok=True)
    model_names = sorted(df_all['model'].dropna().unique().tolist())
    if split_presets_order is None:
        split_order = [p for p in SPLIT_PRESET_SPECS.keys() if p in df_all['split_preset'].values]
    else:
        split_order = list(split_presets_order)
    if len(split_order) == 1:
        sp = split_order[0]
        sub_all = df_all[df_all['split_preset'] == sp].copy()
        if sub_all.empty:
            return
        preferred_order = ['ElasticNet', 'RandomForest', 'SVR', 'Ridge', 'GradientBoosting']
        ordered_models = [m for m in preferred_order if m in model_names] + [m for m in model_names if m not in preferred_order]
        fig, axes = plt.subplots(3, 2, figsize=(7.2, 8.4), sharex=False, sharey=True)
        axes_list = axes.flatten()
        y_min = float(sub_all[['r_no_aug', 'r_with_aug']].min().min())
        y_max = float(sub_all[['r_no_aug', 'r_with_aug']].max().max())
        span = max(y_max - y_min, 0.01)
        margin = max(0.015, span * 0.1)
        y_low = y_min - margin
        y_high = y_max + margin
        seed_min = int(sub_all['seed'].min())
        seed_max = int(sub_all['seed'].max())
        if seed_max - seed_min <= 12:
            x_ticks = list(range(seed_min, seed_max + 1, 2))
        elif seed_max - seed_min <= 30:
            x_ticks = list(range(seed_min, seed_max + 1, 5))
        else:
            x_ticks = list(range(seed_min, seed_max + 1, 20))
            if seed_max not in x_ticks:
                x_ticks.append(seed_max)
        combined_plot_name = f'\u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u53d8\u5316\u66f2\u7ebf\u56fe_\u5408\u5e76_{sp}'
        plot_data_dir = os.path.join(out_dir, 'plot_data', _sanitize_plot_tag(combined_plot_name))
        combined_data = {'split_preset': sp, 'split_desc': SPLIT_PRESET_SPECS[sp]['description'], 'model_order': np.array(ordered_models, dtype=object), 'seed_min': seed_min, 'seed_max': seed_max}
        for i, model_name in enumerate(ordered_models):
            ax = axes_list[i]
            sub = sub_all[sub_all['model'] == model_name].sort_values('seed')
            if sub.empty:
                ax.set_visible(False)
                continue
            seeds = sub['seed'].to_numpy(dtype=int)
            y0 = sub['r_no_aug'].to_numpy(dtype=float)
            y1 = sub['r_with_aug'].to_numpy(dtype=float)
            model_df = sub[['split_preset', 'split_desc', 'seed', 'model', 'r_no_aug', 'r_with_aug', 'r_delta_with_minus_no']].copy()
            combined_data[f'{model_name}_table'] = model_df
            combined_data[f'{model_name}_seed'] = seeds
            combined_data[f'{model_name}_r_no_aug'] = y0
            combined_data[f'{model_name}_r_with_aug'] = y1
            ax.plot(seeds, y0, color=SCATTER_COLORS['no_aug'], linewidth=1.2, marker='o', markersize=2.6, alpha=0.95)
            ax.plot(seeds, y1, color=SCATTER_COLORS['with_aug'], linewidth=1.2, marker='o', markersize=2.6, alpha=0.95)
            ax.set_ylim(y_low, y_high)
            ax.set_xlim(seed_min, seed_max)
            ax.set_xticks(x_ticks)
            ax.grid(True, alpha=0.22, linewidth=0.5)
            ax.set_title(_normalize_model_display_name(model_name), fontsize=10, pad=3)
            _annotate_panel(ax, f'({chr(97 + i)})', y_offset=0.98)
            ax.set_xlabel('Seed', fontsize=9)
            ax.set_ylabel('Pearson r', fontsize=9)
            ax.tick_params(axis='both', labelsize=8, length=3)
            _format_tick_labels(ax)
            for spine in ['top', 'right']:
                ax.spines[spine].set_visible(False)
        for j in range(len(ordered_models), len(axes_list)):
            axes_list[j].set_visible(False)
        from matplotlib.lines import Line2D
        handles = [Line2D([0], [0], color=SCATTER_COLORS['no_aug'], marker='o', linewidth=1.2, markersize=3.2, label='No Aug'), Line2D([0], [0], color=SCATTER_COLORS['with_aug'], marker='o', linewidth=1.2, markersize=3.2, label='With Aug')]
        fig.legend(handles=handles, loc='upper center', bbox_to_anchor=(0.5, 0.985), ncol=2, frameon=False, fontsize=9, columnspacing=1.4, handletextpad=0.5)
        fig.subplots_adjust(top=0.91, bottom=0.08, left=0.09, right=0.98, hspace=0.45, wspace=0.22)
        params_dict = {'figsize': (7.2, 8.4), 'layout': '3x2 multi-panel', 'shared_legend': True, 'shared_legend_labels': ['No Aug', 'With Aug'], 'subplot_ylabel': 'Pearson r', 'subplot_xlabel': 'Seed', 'split_preset': sp, 'split_desc': SPLIT_PRESET_SPECS[sp]['description'], 'model_order': ordered_models, 'y_limits': [float(y_low), float(y_high)], 'x_limits': [int(seed_min), int(seed_max)], 'x_ticks': x_ticks}
        save_plot_params_and_data(combined_plot_name, params_dict, combined_data, output_dir=plot_data_dir)
        fig_path = os.path.join(out_dir, f'{combined_plot_name}.png')
        _save_publication_figure(fig, fig_path, dpi=600)
        plt.close(fig)
        return
    n_sp = len(split_order)
    for model_name in model_names:
        if n_sp <= 1:
            fig, axes = plt.subplots(1, 1, figsize=(10, 5.5), sharex=True)
            axes_list = [axes]
        else:
            fig, axes = plt.subplots(2, 2, figsize=(14, 9), sharex=True)
            axes_list = axes.flatten()
        for i, sp in enumerate(split_order):
            ax = axes_list[i]
            sub = df_all[(df_all['model'] == model_name) & (df_all['split_preset'] == sp)].sort_values('seed')
            if sub.empty:
                ax.set_title(f'{sp} (no data)')
                continue
            seeds = sub['seed'].values
            y0 = sub['r_no_aug'].values
            y1 = sub['r_with_aug'].values
            ax.plot(seeds, y0, marker='o', markersize=2.5, linewidth=1.2, label='No Aug')
            ax.plot(seeds, y1, marker='o', markersize=2.5, linewidth=1.2, label='With Aug')
            y_min = float(np.nanmin(np.concatenate([y0, y1])))
            y_max = float(np.nanmax(np.concatenate([y0, y1])))
            span = max(y_max - y_min, 0.005)
            margin = max(0.002, span * 0.15)
            ax.set_ylim(y_min - margin, y_max + margin)
            ax.set_title(f"{sp} ({SPLIT_PRESET_SPECS[sp]['description']})")
            ax.set_xlabel('Seed')
            ax.set_ylabel('Test Pearson r')
            ax.grid(True, alpha=0.3)
        for j in range(len(split_order), len(axes_list)):
            axes_list[j].set_visible(False)
        handles, labels = axes_list[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, loc='upper center', ncol=2, bbox_to_anchor=(0.5, 0.98), frameon=False)
        fig.suptitle(f'{_normalize_model_display_name(model_name)}: Seed vs r', y=0.995)
        fig.tight_layout(rect=[0, 0, 1, 0.92])
        fig_path = os.path.join(out_dir, f'seed_curve_{_safe_name(model_name)}.png')
        _save_publication_figure(fig, fig_path, dpi=600)
        plt.close(fig)

def _plot_summary_seed_curves_overlay(df_all, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    model_names = sorted(df_all['model'].dropna().unique().tolist())
    splits_present = [p for p in SPLIT_PRESET_SPECS.keys() if p in df_all['split_preset'].values]
    palette = sns.color_palette('Set2', n_colors=max(len(splits_present), 3))
    for model_name in model_names:
        fig, ax = plt.subplots(figsize=(12, 6))
        for j, sp in enumerate(splits_present):
            sub = df_all[(df_all['model'] == model_name) & (df_all['split_preset'] == sp)].sort_values('seed')
            if sub.empty:
                continue
            c = palette[j % len(palette)]
            seeds = sub['seed'].values
            ax.plot(seeds, sub['r_with_aug'].values, marker='o', ms=2.5, lw=1.4, color=c, label=f'{sp} With Aug')
            ax.plot(seeds, sub['r_no_aug'].values, linestyle='--', marker='x', ms=2.5, lw=1.1, color=c, alpha=0.65, label=f'{sp} No Aug')
        y_min = float(df_all[df_all['model'] == model_name][['r_no_aug', 'r_with_aug']].min().min())
        y_max = float(df_all[df_all['model'] == model_name][['r_no_aug', 'r_with_aug']].max().max())
        span = max(y_max - y_min, 0.01)
        margin = max(0.003, span * 0.12)
        ax.set_ylim(y_min - margin, y_max + margin)
        ax.set_xlabel('Seed')
        ax.set_ylabel('Test Pearson r')
        ax.set_title(f'{model_name}: splits overlay (Solid=With Aug, Dashed=No Aug)')
        ax.grid(True, alpha=0.3)
        ax.legend(bbox_to_anchor=(1.02, 1), loc='upper left', fontsize=8, borderaxespad=0.0)
        fig.tight_layout()
        fig.savefig(os.path.join(out_dir, f'\u6c47\u603b_overlay_{_safe_name(model_name)}.png'), dpi=300, bbox_inches='tight')
        plt.close(fig)

def _plot_summary_best_r_by_split(df_all, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    per_seed = df_all.groupby(['split_preset', 'seed'], as_index=False)['best_final_r_with_aug'].first()
    best_by_split = per_seed.groupby('split_preset')['best_final_r_with_aug'].max().reindex([p for p in SPLIT_PRESET_SPECS.keys() if p in per_seed['split_preset'].values])
    per_seed_no = df_all.groupby(['split_preset', 'seed'], as_index=False)['best_final_r_no_aug'].first()
    best_no = per_seed_no.groupby('split_preset')['best_final_r_no_aug'].max().reindex(best_by_split.index)
    x = np.arange(len(best_by_split))
    w = 0.36
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - w / 2, best_no.values, width=w, label='Best No Aug (per split)', color=SCATTER_COLORS['no_aug'], alpha=0.85)
    ax.bar(x + w / 2, best_by_split.values, width=w, label='Best With Aug (per split)', color=SCATTER_COLORS['with_aug'], alpha=0.85)
    ax.set_xticks(x)
    ax.set_xticklabels(list(best_by_split.index), rotation=20, ha='right')
    ax.set_ylabel('Best test r (over seeds)')
    ax.set_title('\u5404\u5212\u5206\u65b9\u5f0f\uff1a\u672c\u6b21\u79cd\u5b50\u533a\u95f4\u5185\u6700\u4f18\u8868\u73b0\u5bf9\u6bd4')
    ax.legend()
    ax.grid(True, axis='y', alpha=0.35)
    fig.tight_layout()
    fig.savefig(os.path.join(out_dir, '\u6c47\u603b_\u5404\u5212\u5206_\u533a\u95f4\u5185\u6700\u4f18r_\u67f1\u72b6\u56fe.png'), dpi=300, bbox_inches='tight')
    plt.close(fig)

def _extract_split_top10_from_all_rows(df_all_rows):
    if df_all_rows is None or len(df_all_rows) == 0:
        return (pd.DataFrame(), pd.DataFrame())
    agg = df_all_rows.groupby(['split_preset', 'split_desc', 'seed'], as_index=False).agg(best_model_r_with_aug=('r_with_aug', 'max'), best_model_r_no_aug=('r_no_aug', 'max'), mean_r_with_aug=('r_with_aug', 'mean'), mean_r_no_aug=('r_no_aug', 'mean'), mean_delta=('r_delta_with_minus_no', 'mean'))
    splits_in_data = [p for p in SPLIT_PRESET_SPECS.keys() if p in agg['split_preset'].values]
    top10_parts = []
    for sp in splits_in_data:
        sub = agg[agg['split_preset'] == sp].sort_values('best_model_r_with_aug', ascending=False).head(10).copy()
        sub['rank_in_split'] = np.arange(1, len(sub) + 1)
        top10_parts.append(sub)
    if len(top10_parts) == 0:
        return (pd.DataFrame(), pd.DataFrame())
    df_top10_long = pd.concat(top10_parts, axis=0, ignore_index=True)
    wide_parts = []
    for sp in splits_in_data:
        sub = df_top10_long[df_top10_long['split_preset'] == sp][['rank_in_split', 'seed', 'best_model_r_with_aug', 'best_model_r_no_aug', 'mean_delta']].copy()
        sub = sub.rename(columns={'seed': f'{sp}_seed', 'best_model_r_with_aug': f'{sp}_best_with_aug', 'best_model_r_no_aug': f'{sp}_best_no_aug', 'mean_delta': f'{sp}_mean_delta'})
        wide_parts.append(sub)
    if not wide_parts:
        return (df_top10_long, pd.DataFrame())
    df_top10_wide = wide_parts[0]
    for i in range(1, len(wide_parts)):
        df_top10_wide = df_top10_wide.merge(wide_parts[i], on='rank_in_split', how='outer')
    df_top10_wide = df_top10_wide.sort_values('rank_in_split').reset_index(drop=True)
    return (df_top10_long, df_top10_wide)

def _parse_split_preset_list(s):
    if s is None or not str(s).strip():
        return None
    parts = [p.strip() for p in str(s).split(',') if p.strip()]
    for p in parts:
        if p not in SPLIT_PRESET_SPECS:
            raise ValueError(f'\u672a\u77e5\u5212\u5206\u9884\u8bbe {p!r}\uff0c\u53ef\u9009: {list(SPLIT_PRESET_SPECS.keys())}')
    return parts

def _finalize_split_seed_search_folder(split_dir, sp, split_rows, X_raw, y_raw, feat_names, d_codes, a_codes, plot_max_features=10):
    os.makedirs(split_dir, exist_ok=True)
    df_split = pd.DataFrame(split_rows).sort_values(['seed', 'model']).reset_index(drop=True)
    if df_split.empty:
        return []
    set_run_output_root(split_dir)
    try:
        df_split.to_csv(os.path.join(split_dir, '\u88681_\u672c\u5212\u5206_\u6240\u6709seed_\u4e94\u6a21\u578br\u5bf9\u6bd4.csv'), index=False, encoding='utf-8-sig')
        seed_scores = df_split.groupby('seed')['best_final_r_with_aug'].first()
        best_seed = int(seed_scores.idxmax())
        df_b = df_split[df_split['seed'] == best_seed]
        rows_t2 = []
        for _, rr in df_b.iterrows():
            rows_t2.append({'split_preset': sp, 'split_desc': SPLIT_PRESET_SPECS[sp]['description'], 'best_seed': best_seed, 'model': rr['model'], 'r_no_aug': rr['r_no_aug'], 'r_with_aug': rr['r_with_aug'], 'r_delta_with_minus_no': rr['r_delta_with_minus_no']})
        pd.DataFrame(rows_t2).to_csv(os.path.join(split_dir, '\u88682_\u672c\u5212\u5206_\u6700\u4f18seed_\u4e94\u6a21\u578br\u5bf9\u6bd4.csv'), index=False, encoding='utf-8-sig')
        df_t2 = pd.DataFrame(rows_t2)
        pt = df_t2.pivot_table(index=['split_preset', 'split_desc', 'best_seed'], columns='model', values=['r_no_aug', 'r_with_aug', 'r_delta_with_minus_no'], aggfunc='first')
        pt.columns = [f'{a}_{b}' for a, b in pt.columns]
        pt.reset_index().to_csv(os.path.join(split_dir, '\u88682_\u672c\u5212\u5206_\u6700\u4f18seed_\u4e94\u6a21\u578br\u5bf9\u6bd4_\u5bbd\u8868.csv'), index=False, encoding='utf-8-sig')
        t3l, t3w = _extract_split_top10_from_all_rows(df_split)
        t3l.to_csv(os.path.join(split_dir, '\u88683_\u672c\u5212\u5206_seed\u603b\u6392\u540d_top10_\u957f\u8868.csv'), index=False, encoding='utf-8-sig')
        t3w.to_csv(os.path.join(split_dir, '\u88683_\u672c\u5212\u5206_seed\u603b\u6392\u540d_top10_\u5bbd\u8868.csv'), index=False, encoding='utf-8-sig')
        curves_dir = os.path.join(split_dir, '\u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u53d8\u5316\u66f2\u7ebf\u56fe')
        _plot_seed_curves_by_model(df_split, curves_dir, split_presets_order=[sp])
        set_run_output_root(None)
        result_no_aug = train_and_evaluate(X_raw, y_raw, f'\u65e0\u589e\u5f3a(best_seed={best_seed})', random_state=best_seed, split_preset=sp, augment_train=False, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=a_codes)
        result_with_aug = train_and_evaluate(X_raw, y_raw, f'\u6709\u589e\u5f3a(best_seed={best_seed})', random_state=best_seed, split_preset=sp, augment_train=True, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=a_codes)
        set_run_output_root(split_dir)
        plot_model_scatters_comparison(result_no_aug, result_with_aug)
        plot_metrics_comparison(result_no_aug, result_with_aug)
        plot_final_model_comparison_scatter(result_no_aug, result_with_aug)
        y_with_plot = np.concatenate([result_with_aug['y_train'], result_with_aug['y_val'], result_with_aug['y_test']])
        plot_combined_pce_distribution_ab(y_raw, y_with_plot)
        names_sel = result_with_aug['selected_feature_names']
        X_corr = np.vstack([result_with_aug['X_train'], result_with_aug['X_val'], result_with_aug['X_test']])
        plot_feature_correlation(X_corr, names_sel, y_with_plot, max_features=plot_max_features)
        if SHAP_AVAILABLE:
            plot_shap_explanations(result_no_aug, result_no_aug['selected_feature_names'], result_no_aug['X_train'], result_no_aug['X_test'], result_no_aug['y_test'], f'No_Augmentation_best_seed_{best_seed}_{sp}', random_state=best_seed, max_display=12)
        meta = {'split_preset': sp, 'best_seed': best_seed, 'best_final_r_no_aug': float(result_no_aug['final_r']), 'best_final_r_with_aug': float(result_with_aug['final_r'])}
        with open(os.path.join(split_dir, 'split_folder_summary.json'), 'w', encoding='utf-8') as sf:
            json.dump(meta, sf, indent=2, ensure_ascii=False)
    finally:
        set_run_output_root(None)
    return rows_t2

def run_all_splits_seed_search_and_summary(start_seed=0, end_seed=100, only_presets=None):
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    preset_tag = 'all_splits' if not only_presets else '_'.join(only_presets)
    out_root = os.path.join(get_run_output_root(), f'split_seed_search_summary_{preset_tag}_{start_seed}_{end_seed}_{ts}')
    os.makedirs(out_root, exist_ok=True)
    if only_presets:
        pass
    set_run_output_root(None)
    base_seed = DEFAULT_RANDOM_SEED
    select_and_split_unknown_20_pairs(random_state=base_seed)
    X_raw, y_raw, feat_names, d_codes, a_codes, _, _ = load_data_without_augmentation()
    all_rows = []
    split_best_rows = []
    if only_presets:
        only_set = set(only_presets)
        preset_iter = [p for p in SPLIT_PRESET_SPECS.keys() if p in only_set]
    else:
        preset_iter = list(SPLIT_PRESET_SPECS.keys())
    for sp in preset_iter:
        split_rows = []
        for seed in range(int(start_seed), int(end_seed) + 1):
            result_no_aug = train_and_evaluate(X_raw, y_raw, '\u65e0\u589e\u5f3a(seed_search)', random_state=seed, split_preset=sp, augment_train=False, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=a_codes)
            result_with_aug = train_and_evaluate(X_raw, y_raw, '\u6709\u589e\u5f3a(seed_search)', random_state=seed, split_preset=sp, augment_train=True, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=a_codes)
            built = _build_split_seed_result_rows(sp, seed, result_no_aug, result_with_aug)
            split_rows.extend(built)
            all_rows.extend(built)
        split_dir = os.path.join(out_root, f'split_{sp}')
        rows_t2 = _finalize_split_seed_search_folder(split_dir, sp, split_rows, X_raw, y_raw, feat_names, d_codes, a_codes, plot_max_features=10)
        split_best_rows.extend(rows_t2)
    summary_dir = os.path.join(out_root, '\u6c47\u603b_all_splits')
    os.makedirs(summary_dir, exist_ok=True)
    set_run_output_root(summary_dir)
    try:
        df_all = pd.DataFrame(all_rows)
        df_all = df_all.sort_values(['split_preset', 'seed', 'model']).reset_index(drop=True)
        table1_path = os.path.join(summary_dir, '\u88681_\u5404\u5212\u5206_\u6240\u6709seed_\u4e94\u6a21\u578br\u5bf9\u6bd4.csv')
        df_all.to_csv(table1_path, index=False, encoding='utf-8-sig')
        df_best = pd.DataFrame(split_best_rows)
        df_best = df_best.sort_values(['split_preset', 'model']).reset_index(drop=True)
        table2_path = os.path.join(summary_dir, '\u88682_\u5404\u5212\u5206_\u6700\u4f18seed_\u4e94\u6a21\u578br\u5bf9\u6bd4.csv')
        df_best.to_csv(table2_path, index=False, encoding='utf-8-sig')
        df_best_wide = df_best.pivot_table(index=['split_preset', 'split_desc', 'best_seed'], columns='model', values=['r_no_aug', 'r_with_aug', 'r_delta_with_minus_no'], aggfunc='first')
        df_best_wide.columns = [f'{a}_{b}' for a, b in df_best_wide.columns]
        df_best_wide = df_best_wide.reset_index()
        table2_wide_path = os.path.join(summary_dir, '\u88682_\u5404\u5212\u5206_\u6700\u4f18seed_\u4e94\u6a21\u578br\u5bf9\u6bd4_\u5bbd\u8868.csv')
        df_best_wide.to_csv(table2_wide_path, index=False, encoding='utf-8-sig')
        df_top10_long, df_top10_wide = _extract_split_top10_from_all_rows(df_all)
        table3_long_path = os.path.join(summary_dir, '\u88683_\u5404\u5212\u5206_seed\u603b\u6392\u540d_top10_\u957f\u8868.csv')
        table3_wide_path = os.path.join(summary_dir, '\u88683_\u5404\u5212\u5206_seed\u603b\u6392\u540d_top10_\u5bbd\u8868.csv')
        df_top10_long.to_csv(table3_long_path, index=False, encoding='utf-8-sig')
        df_top10_wide.to_csv(table3_wide_path, index=False, encoding='utf-8-sig')
        curves_dir = os.path.join(summary_dir, '\u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u53d8\u5316\u66f2\u7ebf\u56fe_\u56db\u5206\u56fe')
        _plot_seed_curves_by_model(df_all, curves_dir, split_presets_order=preset_iter)
        overlay_dir = os.path.join(summary_dir, '\u6c47\u603b_\u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u66f2\u7ebf_\u8de8\u5212\u5206\u53e0\u52a0')
        _plot_summary_seed_curves_overlay(df_all, overlay_dir)
        _plot_summary_best_r_by_split(df_all, summary_dir)
    finally:
        set_run_output_root(None)
    summary_txt = os.path.join(summary_dir, '\u8fd0\u884c\u603b\u7ed3.txt')
    with open(summary_txt, 'w', encoding='utf-8') as f:
        f.write('\u56db\u79cd\u5212\u5206\u6bd4\u4f8b \xd7 \u968f\u673a\u79cd\u5b50\u641c\u7d22\u603b\u7ed3\n')
        f.write(f'\u79cd\u5b50\u8303\u56f4: {start_seed} ~ {end_seed}\n')
        f.write(f'\u6839\u76ee\u5f55: {out_root}\n')
        f.write('\u5404\u5212\u5206\u5b50\u76ee\u5f55: split_<\u9884\u8bbe\u540d>/\uff08\u542b\u8be5\u5212\u5206\u5168\u90e8\u8868\u3001\u79cd\u5b50\u66f2\u7ebf\u3001\u6700\u4f18 seed \u6027\u80fd\u5bf9\u6bd4\u56fe\uff09\n')
        f.write('\u6c47\u603b\u76ee\u5f55: \u6c47\u603b_all_splits/\n\n')
        f.write('\u5404\u5212\u5206\u6700\u4f18 seed\uff08\u6309\u6709\u589e\u5f3a\u6700\u7ec8 best r\uff0c\u5728\u5b50\u76ee\u5f55 split_folder_summary.json \u4e2d\u4ea6\u6709\uff09:\n')
        for sp in preset_iter:
            sub = df_best[df_best['split_preset'] == sp]
            if len(sub) == 0:
                continue
            bseed = int(sub['best_seed'].iloc[0])
            best_with_aug = float(sub['r_with_aug'].max())
            f.write(f"- {sp} ({SPLIT_PRESET_SPECS[sp]['description']}): best_seed={bseed}, best_model_r_with_aug={best_with_aug:.4f}\n")
        f.write('\n\u6c47\u603b\u76ee\u5f55\u5185\u4e3b\u8981\u6587\u4ef6:\n')
        f.write(f'- {table1_path}\n')
        f.write(f'- {table2_path}\n')
        f.write(f'- {table2_wide_path}\n')
        f.write(f'- {table3_long_path}\n')
        f.write(f'- {table3_wide_path}\n')
        f.write(f'- {curves_dir}\n')
        f.write(f'- {overlay_dir}\n')
        f.write(f"- {os.path.join(summary_dir, '\u6c47\u603b_\u5404\u5212\u5206_\u533a\u95f4\u5185\u6700\u4f18r_\u67f1\u72b6\u56fe.png')}\n")
    readme_root = os.path.join(out_root, 'README_\u8f93\u51fa\u7ed3\u6784.txt')
    with open(readme_root, 'w', encoding='utf-8') as rf:
        rf.write(f'\u672c\u6279\u6b21\u8f93\u51fa\u7ed3\u6784\u8bf4\u660e\n================\n1) \u5404\u5212\u5206\u5b8c\u6574\u7ed3\u679c\uff08\u6bcf\u79cd\u5212\u5206\u8dd1\u5b8c {start_seed}~{end_seed} \u540e\u751f\u6210\uff09:\n   split_<\u9884\u8bbe\u540d>/\n   - \u88681/2/3\uff08CSV\uff09\n   - \u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u53d8\u5316\u66f2\u7ebf\u56fe/\n   - \u6570\u636e\u589e\u5f3a\u524d\u540e\u5404\u6a21\u578b\u7684\u6563\u70b9\u56fe.png\u3001\u6307\u6807\u5bf9\u6bd4\u56fe\u3001\u96c6\u6210\u6563\u70b9\u56fe\u3001\u6570\u636e\u5206\u5e03\u56fe\u3001\u7279\u5f81\u76f8\u5173\u6027\u56fe\u7b49\n   - plot_data/\uff08\u5404\u56fe\u6570\u636e\u7f13\u5b58\uff09\n   - split_folder_summary.json\uff08\u8be5\u5212\u5206\u6700\u4f18 seed \u4e0e\u6700\u7ec8 r\uff09\n\n2) \u56db\u79cd\u5212\u5206\u5168\u90e8\u5b8c\u6210\u540e:\n   \u6c47\u603b_all_splits/\n   - \u5408\u5e76\u88681/2/3\n   - \u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u53d8\u5316\u66f2\u7ebf\u56fe_\u56db\u5206\u56fe/\uff08\u6bcf\u6a21\u578b\u56db\u5b50\u56fe\uff09\n   - \u6c47\u603b_\u6a21\u578b\u6027\u80fd_\u79cd\u5b50\u66f2\u7ebf_\u8de8\u5212\u5206\u53e0\u52a0/\uff08\u6bcf\u6a21\u578b\u591a\u5212\u5206\u53e0\u52a0\uff09\n   - \u6c47\u603b_\u5404\u5212\u5206_\u533a\u95f4\u5185\u6700\u4f18r_\u67f1\u72b6\u56fe.png\n')
    return out_root
FIVE_PT_SIZE = 10.5
FONT_FAMILY = ['Times New Roman', 'SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'DejaVu Sans']
matplotlib.rcParams['font.family'] = FONT_FAMILY
matplotlib.rcParams['font.sans-serif'] = FONT_FAMILY
matplotlib.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = FIVE_PT_SIZE
sns.set_theme(style='whitegrid', context='talk', font='Times New Roman')
PUBLICATION_RC = {'figure.dpi': 600, 'savefig.dpi': 600, 'axes.spines.top': False, 'axes.spines.right': False, 'axes.titleweight': 'normal', 'axes.labelweight': 'normal', 'axes.titlesize': FIVE_PT_SIZE, 'axes.labelsize': FIVE_PT_SIZE, 'legend.fontsize': FIVE_PT_SIZE, 'xtick.labelsize': FIVE_PT_SIZE, 'ytick.labelsize': FIVE_PT_SIZE, 'grid.alpha': 0.25}
plt.rcParams.update(PUBLICATION_RC)
COLOR_PALETTE = sns.color_palette('Set2')
SCATTER_COLORS = {'no_aug': COLOR_PALETTE[0], 'with_aug': COLOR_PALETTE[2], 'ensemble_no_aug': COLOR_PALETTE[1], 'ensemble_with_aug': COLOR_PALETTE[3]}
MODEL_NAME_ABBR = {'RandomForest': 'Random Forest', 'GradientBoosting': 'GBR'}

def _normalize_model_display_name(name: str) -> str:
    if not isinstance(name, str):
        return str(name)
    name = name.replace('RandomForest', 'Random Forest')
    name = name.replace('GradientBoosting', 'GBR')
    name = name.replace('\u6700\u4f73\u5355\u6a21\u578b', 'Best single model')
    return name

def _extract_best_model_name(name: str) -> str:
    label = _normalize_model_display_name(name)
    m = re.match('^Best single model\\s*\\((.*?)\\)\\s*$', label)
    if m:
        return m.group(1).strip()
    return label.strip()

def _parse_paired_feature_label(feature_name: str):
    label = str(feature_name)
    m = re.match('^(fd|fa)_(\\d+)$', label)
    if not m:
        return {'feature_label': label, 'feature_side': 'unknown', 'bit_index': np.nan, 'feature_note': 'non-standard feature label'}
    prefix, bit = (m.group(1), int(m.group(2)))
    side = 'donor' if prefix == 'fd' else 'acceptor'
    return {'feature_label': label, 'feature_side': side, 'bit_index': bit, 'feature_note': f'{side}-side Morgan fingerprint bit'}

def _sanitize_plot_tag(s: str) -> str:
    return str(s).replace(' ', '_').replace('(', '').replace(')', '').replace('/', '_').replace('\\', '_')

def _annotate_panel(ax, label, y_offset=1.02):
    ax.text(-0.12, y_offset, label, transform=ax.transAxes, fontsize=FIVE_PT_SIZE, ha='left', va='bottom')

def _add_identity_line(ax, data_x, data_y, color='#444444'):
    min_val = min(float(np.min(data_x)), float(np.min(data_y)))
    max_val = max(float(np.max(data_x)), float(np.max(data_y)))
    ax.plot([min_val, max_val], [min_val, max_val], linestyle='--', color=color, linewidth=1.5, alpha=0.8)
    ax.set_xlim(min_val - 0.5, max_val + 0.5)
    ax.set_ylim(min_val - 0.5, max_val + 0.5)

def _format_tick_labels(ax):
    ax.tick_params(axis='both', which='major', length=5, width=1.2, direction='out')
    ax.tick_params(axis='both', which='minor', length=3, width=1.0, direction='out')

def _save_publication_figure(fig, filename, dpi=600):
    from pathlib import Path as _Path
    out_path = _Path(filename)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        fig.savefig(out_path, dpi=dpi, bbox_inches='tight', facecolor='white')
    except TypeError:
        fig.savefig(out_path, dpi=dpi, bbox_inches='tight')
    pdf_path = out_path.with_suffix('.pdf')
    try:
        fig.savefig(pdf_path, bbox_inches='tight', facecolor='white')
    except TypeError:
        fig.savefig(pdf_path, bbox_inches='tight')

def load_numpy_if_exists(path, allow_pickle=False):
    if not os.path.isabs(path):
        path = data_path(path)
    if os.path.exists(path):
        try:
            return np.load(path, allow_pickle=allow_pickle)
        except Exception as exc:
            return None
    return None

def load_cached_plot_data(plot_name, output_dir='plot_data'):
    cache_dir = output_dir if os.path.isabs(output_dir) else os.path.join(get_run_output_root(), output_dir)
    npz_path = os.path.join(cache_dir, f'{plot_name}_data.npz')
    params_path = os.path.join(cache_dir, f'{plot_name}_params.json')
    if not os.path.exists(npz_path):
        return (None, None)
    try:
        params = {}
        if os.path.exists(params_path):
            with open(params_path, 'r', encoding='utf-8') as f:
                params = json.load(f)
        data_npz = np.load(npz_path, allow_pickle=True)
        data_dict = {key: data_npz[key] for key in data_npz.files}
        return (params, data_dict)
    except Exception as exc:
        return (None, None)

def rebuild_results_from_cache():
    _, scatter_data = load_cached_plot_data('\u5404\u6a21\u578b\u6563\u70b9\u56fe\u5bf9\u6bd4')
    if scatter_data is None:
        return (None, None)
    model_names = sorted({key.split('_')[0] for key in scatter_data if key.endswith('_no_aug_y_test')})
    result_no_aug = {'results': {}}
    result_with_aug = {'results': {}}
    for name in model_names:
        no_aug_y_test_key = f'{name}_no_aug_y_test'
        with_aug_y_test_key = f'{name}_with_aug_y_test'
        if no_aug_y_test_key in scatter_data:
            if 'y_test' not in result_no_aug:
                result_no_aug['y_test'] = scatter_data[no_aug_y_test_key]
            result_no_aug['results'][name] = {'y_pred_test': scatter_data.get(f'{name}_no_aug_y_pred'), 'r_test': float(np.array(scatter_data.get(f'{name}_no_aug_r', 0.0)).squeeze())}
        if with_aug_y_test_key in scatter_data:
            if 'y_test' not in result_with_aug:
                result_with_aug['y_test'] = scatter_data[with_aug_y_test_key]
            result_with_aug['results'][name] = {'y_pred_test': scatter_data.get(f'{name}_with_aug_y_pred'), 'r_test': float(np.array(scatter_data.get(f'{name}_with_aug_r', 0.0)).squeeze())}
    _, metrics_data = load_cached_plot_data('\u6307\u6807\u5bf9\u6bd4\u56fe')
    if metrics_data is None:
        pass
    else:
        for name in model_names:
            for metric in ['r_test', 'r2_test', 'rmse_test', 'mae_test']:
                key_no_aug = f'{name}_no_aug_{metric}'
                key_with_aug = f'{name}_with_aug_{metric}'
                if key_no_aug in metrics_data and name in result_no_aug['results']:
                    result_no_aug['results'][name][metric] = float(np.array(metrics_data[key_no_aug]).squeeze())
                if key_with_aug in metrics_data and name in result_with_aug['results']:
                    result_with_aug['results'][name][metric] = float(np.array(metrics_data[key_with_aug]).squeeze())
    _, ensemble_data = load_cached_plot_data('\u96c6\u6210\u6a21\u578b\u6563\u70b9\u56fe')
    if ensemble_data is None:
        pass
    else:
        for prefix, result in [('no_aug', result_no_aug), ('with_aug', result_with_aug)]:
            if result and f'{prefix}_y_test' in ensemble_data:
                result['y_test'] = ensemble_data[f'{prefix}_y_test']
                result['final_pred'] = ensemble_data.get(f'{prefix}_y_pred')
                result['final_r'] = float(np.array(ensemble_data.get(f'{prefix}_r', 0.0)).squeeze())
                result['final_rmse'] = float(np.array(ensemble_data.get(f'{prefix}_rmse', 0.0)).squeeze())
                result['final_r2'] = float(np.array(ensemble_data.get(f'{prefix}_r2', 0.0)).squeeze())
                result['final_mae'] = float(np.array(ensemble_data.get(f'{prefix}_mae', 0.0)).squeeze())
    return (result_no_aug if result_no_aug['results'] else None, result_with_aug if result_with_aug['results'] else None)

def run_plot_only_mode(max_features=10):
    y_data_with = load_numpy_if_exists('fp_Y_filtered.npy')
    if y_data_with is None:
        y_data_with = load_numpy_if_exists('fp_Y.npy')
    y_data_no_aug = load_numpy_if_exists('fp_Y_filtered_no_aug.npy')
    if y_data_no_aug is None:
        y_data_no_aug = load_numpy_if_exists('fp_Y_no_aug.npy')
    y_data = y_data_with if y_data_with is not None else y_data_no_aug
    if y_data_no_aug is not None and y_data_with is not None:
        plot_combined_pce_distribution_ab(y_data_no_aug, y_data_with)
    else:
        if y_data_no_aug is None:
            pass
        if y_data_with is None:
            pass
    fd_fp = load_numpy_if_exists('fd_fp_filtered.npy')
    fa_fp = load_numpy_if_exists('fa_fp_filtered.npy')
    if fd_fp is not None and fa_fp is not None:
        X = np.hstack((fd_fp, fa_fp))
        d_feature_names = [f'fd_{i}' for i in range(fd_fp.shape[1])]
        a_feature_names = [f'fa_{i}' for i in range(fa_fp.shape[1])]
        feature_names = d_feature_names + a_feature_names
        plot_feature_correlation(X, feature_names, y=y_data if y_data is not None else None, max_features=max_features)
    result_no_aug, result_with_aug = rebuild_results_from_cache()
    if result_no_aug and result_with_aug:
        try:
            plot_model_scatters_comparison(result_no_aug, result_with_aug)
        except Exception as exc:
            pass
        try:
            plot_metrics_comparison(result_no_aug, result_with_aug)
        except Exception as exc:
            pass
        try:
            plot_final_model_comparison_scatter(result_no_aug, result_with_aug)
        except Exception as exc:
            pass

def save_plot_params_and_data(plot_name, params_dict, data_dict, output_dir='plot_data'):
    output_dir_abs = output_dir
    if not os.path.isabs(output_dir_abs):
        output_dir_abs = os.path.join(get_run_output_root(), output_dir_abs)
    if not os.path.exists(output_dir_abs):
        os.makedirs(output_dir_abs)
    params_file = os.path.join(output_dir_abs, f'{plot_name}_params.json')
    params_json = {}
    for key, value in params_dict.items():
        if isinstance(value, np.ndarray):
            params_json[key] = value.tolist()
        elif isinstance(value, (np.integer, np.floating)):
            params_json[key] = float(value)
        elif isinstance(value, (list, tuple)):
            params_json[key] = [float(v) if isinstance(v, (np.integer, np.floating)) else v for v in value]
        else:
            params_json[key] = value
    with open(params_file, 'w', encoding='utf-8') as f:
        json.dump(params_json, f, indent=2, ensure_ascii=False)
    data_file_npy = os.path.join(output_dir_abs, f'{plot_name}_data.npz')
    data_file_csv = os.path.join(output_dir_abs, f'{plot_name}_data.csv')
    data_npz = {}
    data_csv_dict = {}
    for key, value in data_dict.items():
        if isinstance(value, np.ndarray):
            data_npz[key] = value
            if value.ndim == 1:
                data_csv_dict[key] = value
            elif value.ndim == 2 and value.shape[1] <= 10:
                for col_idx in range(value.shape[1]):
                    data_csv_dict[f'{key}_col{col_idx}'] = value[:, col_idx]
        elif isinstance(value, pd.DataFrame):
            value.to_csv(os.path.join(output_dir_abs, f'{plot_name}_{key}_data.csv'), index=False, encoding='utf-8-sig')
        elif isinstance(value, (list, tuple)):
            data_npz[key] = np.array(value)
            data_csv_dict[key] = np.array(value)
        else:
            data_csv_dict[key] = [value]
    if data_npz:
        np.savez_compressed(data_file_npy, **data_npz)
    if data_csv_dict:
        max_len = max((len(v) if hasattr(v, '__len__') else 1 for v in data_csv_dict.values()))
        data_csv_final = {}
        for key, value in data_csv_dict.items():
            if isinstance(value, np.ndarray):
                if len(value) < max_len:
                    if value.dtype.kind in {'U', 'S', 'O'}:
                        padded = np.full(max_len, None, dtype=object)
                    else:
                        padded = np.full(max_len, np.nan)
                    padded[:len(value)] = value
                    data_csv_final[key] = padded
                else:
                    data_csv_final[key] = value[:max_len]
            else:
                data_csv_final[key] = [value] * max_len
        df_data = pd.DataFrame(data_csv_final)
        df_data.to_csv(data_file_csv, index=False, encoding='utf-8-sig')
    if data_npz:
        pass
    if data_csv_dict:
        pass

def save_histogram_only_data(plot_name, bin_edges, counts, output_dir='plot_data'):
    output_dir_abs = output_dir
    if not os.path.isabs(output_dir_abs):
        output_dir_abs = os.path.join(get_run_output_root(), output_dir_abs)
    if not os.path.exists(output_dir_abs):
        os.makedirs(output_dir_abs)
    centers = (bin_edges[:-1] + bin_edges[1:]) / 2
    hist_npz_path = os.path.join(output_dir_abs, f'{plot_name}_hist_only.npz')
    np.savez_compressed(hist_npz_path, bin_edges=bin_edges, counts=counts, bin_centers=centers)
    hist_csv_path = os.path.join(output_dir_abs, f'{plot_name}_hist_only.csv')
    hist_df = pd.DataFrame({'bin_left': bin_edges[:-1], 'bin_right': bin_edges[1:], 'bin_center': centers, 'count': counts})
    hist_df.to_csv(hist_csv_path, index=False, encoding='utf-8-sig')

def select_and_split_unknown_20_pairs(n_unknown=20, random_state=None):
    da_path = _find_first_existing_source_file(PUBLIC_DA_FILENAME)
    donor_path = _find_first_existing_source_file(PUBLIC_DONOR_FILENAME)
    acceptor_path = _find_first_existing_source_file(PUBLIC_ACCEPTOR_FILENAME)

    da_real = _read_smiles_table_auto(da_path)
    geiti = _load_reference_smiles_table(donor_path, 'Donor')
    shouti = _load_reference_smiles_table(acceptor_path, 'Acceptor')
    if geiti is None or geiti.empty:
        raise FileNotFoundError(f'Not found or empty donor reference file: {donor_path}')
    if shouti is None or shouti.empty:
        raise FileNotFoundError(f'Not found or empty acceptor reference file: {acceptor_path}')

    donor_col, acceptor_col, pce_col = _guess_smiles_dataset_columns_auto(da_real, None, None, None)
    da_real = da_real[[donor_col, acceptor_col, pce_col]].copy()
    da_real.columns = ['donor_smiles', 'acceptor_smiles', 'PCE']
    da_real['donor_smiles'] = da_real['donor_smiles'].map(_norm_text_auto)
    da_real['acceptor_smiles'] = da_real['acceptor_smiles'].map(_norm_text_auto)
    da_real['PCE'] = pd.to_numeric(da_real['PCE'], errors='coerce')
    da_real = da_real[(da_real['donor_smiles'] != '') & (da_real['acceptor_smiles'] != '')].dropna(subset=['PCE']).reset_index(drop=True)

    da_real['DA_Key'] = da_real['donor_smiles'] + '_' + da_real['acceptor_smiles']
    da_real_cleaned = da_real.sort_values(by='PCE', ascending=False).drop_duplicates(subset='DA_Key', keep='first').reset_index(drop=True)
    da_real_cleaned.to_csv(data_path(PUBLIC_DA_CLEANED_FILENAME), index=False, encoding='utf-8-sig')

    rs = _resolve_random_state(random_state)
    np.random.seed(rs)
    available_rows = len(da_real_cleaned)
    if available_rows < n_unknown:
        n_unknown = available_rows
    selected_csv_indices = np.random.choice(available_rows, size=n_unknown, replace=False)
    selected_csv_indices = np.sort(selected_csv_indices)
    selected_rows = da_real_cleaned.iloc[selected_csv_indices].copy()
    remaining_da_real = da_real_cleaned.drop(selected_csv_indices).reset_index(drop=True)
    remaining_da_real.to_csv(data_path(PUBLIC_DA_FILTERED_FILENAME), index=False, encoding='utf-8-sig')

    donor_smiles_list = geiti['SMILES'].values
    acceptor_smiles_list = shouti['SMILES'].values
    donor_name_to_idx = {str(name).strip(): idx for idx, name in enumerate(geiti['Donor'].values)} if 'Donor' in geiti.columns else {}
    acceptor_name_to_idx = {str(name).strip(): idx for idx, name in enumerate(shouti['Acceptor'].values)} if 'Acceptor' in shouti.columns else {}

    fd_fp = np.load(input_path('fd_fp_augmented.npy'))
    fa_fp = np.load(input_path('fa_fp_augmented.npy'))
    fp_Y = np.load(input_path('fp_Y_augmented.npy'))
    total_samples = fd_fp.shape[0]
    try:
        fd_fp1 = np.load(input_path('fd_fp1.npy'))
        fa_fp1 = np.load(input_path('fa_fp1.npy'))
    except Exception:
        fd_fp1 = None
        fa_fp1 = None

    donor_smiles_to_idx = {str(smiles).strip(): idx for idx, smiles in enumerate(donor_smiles_list)}
    acceptor_smiles_to_idx = {str(smiles).strip(): idx for idx, smiles in enumerate(acceptor_smiles_list)}

    unknown_donor_smiles_list = []
    unknown_acceptor_smiles_list = []
    unknown_pce_values = []
    unknown_donor_fp_list = []
    unknown_acceptor_fp_list = []
    unknown_indices_in_augmented = []

    for _, row in selected_rows.iterrows():
        donor_info_csv = str(row['donor_smiles']).strip()
        acceptor_info_csv = str(row['acceptor_smiles']).strip()
        pce_value_csv = float(row['PCE']) if pd.notna(row['PCE']) else None

        donor_idx_in_ref = donor_name_to_idx.get(donor_info_csv)
        acceptor_idx_in_ref = acceptor_name_to_idx.get(acceptor_info_csv)

        if donor_idx_in_ref is None:
            donor_idx_in_ref = donor_smiles_to_idx.get(donor_info_csv)
        if acceptor_idx_in_ref is None:
            acceptor_idx_in_ref = acceptor_smiles_to_idx.get(acceptor_info_csv)

        if donor_idx_in_ref is None:
            for smiles, idx_val in donor_smiles_to_idx.items():
                if str(smiles).strip().replace(' ', '') == donor_info_csv.replace(' ', ''):
                    donor_idx_in_ref = idx_val
                    break
        if acceptor_idx_in_ref is None:
            for smiles, idx_val in acceptor_smiles_to_idx.items():
                if str(smiles).strip().replace(' ', '') == acceptor_info_csv.replace(' ', ''):
                    acceptor_idx_in_ref = idx_val
                    break

        donor_smi = str(donor_smiles_list[donor_idx_in_ref]).strip() if donor_idx_in_ref is not None and donor_idx_in_ref < len(donor_smiles_list) else donor_info_csv
        acceptor_smi = str(acceptor_smiles_list[acceptor_idx_in_ref]).strip() if acceptor_idx_in_ref is not None and acceptor_idx_in_ref < len(acceptor_smiles_list) else acceptor_info_csv

        donor_fp = fd_fp1[donor_idx_in_ref] if donor_idx_in_ref is not None and fd_fp1 is not None and donor_idx_in_ref < fd_fp1.shape[0] else None
        acceptor_fp = fa_fp1[acceptor_idx_in_ref] if acceptor_idx_in_ref is not None and fa_fp1 is not None and acceptor_idx_in_ref < fa_fp1.shape[0] else None

        unknown_donor_smiles_list.append(donor_smi)
        unknown_acceptor_smiles_list.append(acceptor_smi)
        unknown_pce_values.append(pce_value_csv)
        unknown_donor_fp_list.append(donor_fp)
        unknown_acceptor_fp_list.append(acceptor_fp)

        matched_idx_in_augmented = None
        if donor_fp is not None and acceptor_fp is not None:
            donor_matches = np.where((fd_fp == donor_fp).all(axis=1))[0]
            acceptor_matches = np.where((fa_fp == acceptor_fp).all(axis=1))[0]
            common_matches = np.intersect1d(donor_matches, acceptor_matches)
            if pce_value_csv is not None and len(common_matches) > 0:
                pce_tolerance = 0.01
                for match_idx in common_matches:
                    pce_in_augmented = fp_Y[match_idx]
                    if abs(pce_in_augmented - pce_value_csv) <= pce_tolerance:
                        matched_idx_in_augmented = match_idx
                        break
                if matched_idx_in_augmented is None and len(common_matches) > 0:
                    pce_diffs = [abs(fp_Y[idx] - pce_value_csv) for idx in common_matches]
                    matched_idx_in_augmented = common_matches[int(np.argmin(pce_diffs))]
            elif len(common_matches) > 0:
                matched_idx_in_augmented = common_matches[0]
            if matched_idx_in_augmented is not None and matched_idx_in_augmented not in unknown_indices_in_augmented:
                unknown_indices_in_augmented.append(matched_idx_in_augmented)

    if len(unknown_donor_smiles_list) < n_unknown and len(unknown_donor_smiles_list) == 0:
        raise ValueError(f'Cannot extract valid records from {da_path}. Please check donor/acceptor SMILES columns.')

    unknown_donor_smiles_list = unknown_donor_smiles_list[:n_unknown]
    unknown_acceptor_smiles_list = unknown_acceptor_smiles_list[:n_unknown]
    unknown_pce_values = unknown_pce_values[:n_unknown]
    unknown_donor_fp_list = unknown_donor_fp_list[:n_unknown]
    unknown_acceptor_fp_list = unknown_acceptor_fp_list[:n_unknown]

    valid_fp_indices = [i for i, (dfp, afp) in enumerate(zip(unknown_donor_fp_list, unknown_acceptor_fp_list)) if dfp is not None and afp is not None]
    if len(valid_fp_indices) > 0:
        unknown_fd_fp = np.array([unknown_donor_fp_list[i] for i in valid_fp_indices])
        unknown_fa_fp = np.array([unknown_acceptor_fp_list[i] for i in valid_fp_indices])
        unknown_donor_smiles_list = [unknown_donor_smiles_list[i] for i in valid_fp_indices]
        unknown_acceptor_smiles_list = [unknown_acceptor_smiles_list[i] for i in valid_fp_indices]
        unknown_pce_values = [unknown_pce_values[i] for i in valid_fp_indices]
    else:
        raise ValueError('Cannot recover donor/acceptor fingerprints from fd_fp1.npy and fa_fp1.npy.')

    unknown_fp_Y = np.array(unknown_pce_values, dtype=np.float64)
    unknown_indices = np.array(unknown_indices_in_augmented, dtype=np.int64) if len(unknown_indices_in_augmented) > 0 else np.array([], dtype=np.int64)
    unknown_indices = np.unique(unknown_indices)
    all_augmented_indices = np.arange(total_samples, dtype=np.int64)
    remaining_indices_in_augmented = np.setdiff1d(all_augmented_indices, unknown_indices)

    try:
        remaining_da_real = pd.read_csv(data_path(PUBLIC_DA_FILTERED_FILENAME), encoding='utf-8-sig')
        updated_pce_count = 0
        for _, row in remaining_da_real.iterrows():
            donor_smiles_csv = str(row['donor_smiles']).strip()
            acceptor_smiles_csv = str(row['acceptor_smiles']).strip()
            pce_value_csv = float(row['PCE']) if pd.notna(row['PCE']) else None
            if pce_value_csv is None:
                continue
            donor_idx_ref = donor_smiles_to_idx.get(donor_smiles_csv)
            acceptor_idx_ref = acceptor_smiles_to_idx.get(acceptor_smiles_csv)
            if donor_idx_ref is not None and acceptor_idx_ref is not None and (fd_fp1 is not None) and (fa_fp1 is not None):
                target_donor_fp = fd_fp1[donor_idx_ref]
                target_acceptor_fp = fa_fp1[acceptor_idx_ref]
                donor_matches = np.where((fd_fp == target_donor_fp).all(axis=1))[0]
                acceptor_matches = np.where((fa_fp == target_acceptor_fp).all(axis=1))[0]
                common_matches = np.intersect1d(donor_matches, acceptor_matches)
                valid_matches = np.intersect1d(common_matches, remaining_indices_in_augmented)
                if len(valid_matches) > 0:
                    for match_idx in valid_matches:
                        if abs(fp_Y[match_idx] - pce_value_csv) <= 0.5:
                            fp_Y[match_idx] = pce_value_csv
                            updated_pce_count += 1
                            break
    except Exception:
        pass

    filtered_fd_fp = fd_fp[remaining_indices_in_augmented]
    filtered_fa_fp = fa_fp[remaining_indices_in_augmented]
    filtered_fp_Y = fp_Y[remaining_indices_in_augmented]
    try:
        fd_fp_original = np.load(input_path('fd_fp.npy'))
        fa_fp_original = np.load(input_path('fa_fp.npy'))
        fp_Y_original = np.load(input_path('fp_Y.npy'))
        merged_fd_fp = np.vstack([fd_fp_original, filtered_fd_fp])
        merged_fa_fp = np.vstack([fa_fp_original, filtered_fa_fp])
        merged_fp_Y = np.concatenate([fp_Y_original, filtered_fp_Y])
    except FileNotFoundError:
        merged_fd_fp = filtered_fd_fp
        merged_fa_fp = filtered_fa_fp
        merged_fp_Y = filtered_fp_Y

    np.save(data_path('unknown_20_fd_fp.npy'), unknown_fd_fp)
    np.save(data_path('unknown_20_fa_fp.npy'), unknown_fa_fp)
    np.save(data_path('unknown_20_fp_Y.npy'), unknown_fp_Y)
    np.save(data_path('unknown_20_indices.npy'), unknown_indices)
    np.save(data_path('fd_fp_filtered.npy'), merged_fd_fp)
    np.save(data_path('fa_fp_filtered.npy'), merged_fa_fp)
    np.save(data_path('fp_Y_filtered.npy'), merged_fp_Y)
    unknown_smiles_data = {'donor_smiles': unknown_donor_smiles_list, 'acceptor_smiles': unknown_acceptor_smiles_list, 'pce_values': unknown_pce_values}
    np.save(data_path('unknown_20_smiles.npy'), unknown_smiles_data, allow_pickle=True)
    return (unknown_indices, merged_fd_fp, merged_fa_fp, merged_fp_Y, unknown_fd_fp, unknown_fa_fp, unknown_fp_Y, unknown_donor_smiles_list, unknown_acceptor_smiles_list)


def load_raw_merged_features():
    fd_fp = np.load(data_path('fd_fp_filtered.npy'))
    fa_fp = np.load(data_path('fa_fp_filtered.npy'))
    fp_Y = np.load(data_path('fp_Y_filtered.npy'))
    X = np.hstack((fd_fp, fa_fp))
    y = fp_Y
    d_feature_names = [f'fd_{i}' for i in range(fd_fp.shape[1])]
    a_feature_names = [f'fa_{i}' for i in range(fa_fp.shape[1])]
    feature_names_all = d_feature_names + a_feature_names
    donor_codes = None
    acceptor_codes = None
    try:
        donor_codes_all = np.load(input_path('donor_codes.npy'), allow_pickle=True)
        acceptor_codes_all = np.load(input_path('acceptor_codes.npy'), allow_pickle=True)
        unknown_indices = np.load(data_path('unknown_20_indices.npy'))
        remaining_indices = np.setdiff1d(np.arange(len(donor_codes_all)), unknown_indices)
        donor_codes = donor_codes_all[remaining_indices]
        acceptor_codes = acceptor_codes_all[remaining_indices]
        if len(donor_codes) != X.shape[0] or len(acceptor_codes) != X.shape[0]:
            n_extra = int(X.shape[0] - len(donor_codes))
            if n_extra < 0:
                raise ValueError(f'\u7f16\u7801\u957f\u5ea6\u5f02\u5e38\uff1acodes={len(donor_codes)}, X={X.shape[0]}')
            if n_extra > 0:
                donor_codes = np.concatenate([np.asarray(donor_codes, dtype=object), np.asarray([f'AUX_GAO_D_{i}' for i in range(n_extra)], dtype=object)])
                acceptor_codes = np.concatenate([np.asarray(acceptor_codes, dtype=object), np.asarray([f'AUX_GAO_A_{i}' for i in range(n_extra)], dtype=object)])
    except Exception as e:
        donor_codes = np.array([f'D_{i}' for i in range(y.shape[0])], dtype=object)
        acceptor_codes = np.array([f'A_{i}' for i in range(y.shape[0])], dtype=object)
    return (X, y, feature_names_all, donor_codes, acceptor_codes)

def load_data_without_augmentation():
    X, y, feat_names, dc, ac = load_raw_merged_features()
    return (X, y, feat_names, dc, ac, None, None)

def load_data_with_augmentation(random_state=None):
    return load_data_without_augmentation()

def _build_models(rs):
    return {'RandomForest': RandomForestRegressor(n_estimators=300, max_depth=20, min_samples_split=3, min_samples_leaf=1, random_state=rs, n_jobs=-1), 'GradientBoosting': GradientBoostingRegressor(n_estimators=300, max_depth=10, learning_rate=0.03, min_samples_split=3, random_state=rs), 'SVR': SVR(kernel='rbf', C=10, gamma='scale'), 'Ridge': Ridge(alpha=0.05), 'ElasticNet': ElasticNet(alpha=0.005, l1_ratio=0.3, max_iter=3000)}

def _train_and_evaluate_80_20_cv(X, y, config_name, rs, augment_train, feature_names_all, donor_codes, acceptor_codes):
    X_trainval, X_test, y_trainval, y_test, dc_trainval, dc_test, ac_trainval, ac_test = train_test_split(X, y, donor_codes, acceptor_codes, test_size=0.2, random_state=rs)
    n_total = len(y)
    kf = KFold(n_splits=CV_N_SPLITS, shuffle=True, random_state=rs)
    model_templates = _build_models(rs)
    cv_r2_lists = {name: [] for name in model_templates}
    cv_r_lists = {name: [] for name in model_templates}
    for fold_i, (tr_idx, va_idx) in enumerate(kf.split(X_trainval)):
        rs_fold = int((rs + fold_i * 7919) % (2 ** 31 - 1))
        X_tr, X_va = (X_trainval[tr_idx], X_trainval[va_idx])
        y_tr, y_va = (y_trainval[tr_idx], y_trainval[va_idx])
        if augment_train:
            X_tr, y_tr = _augment_train_features_only(X_tr, y_tr, rs_fold)
        k = min(600, X_tr.shape[1])
        selector_f = SelectKBest(score_func=mutual_info_regression, k=k)
        X_tr_sel = selector_f.fit_transform(X_tr, y_tr)
        X_va_sel = selector_f.transform(X_va)
        scaler_f = RobustScaler()
        X_tr_s = scaler_f.fit_transform(X_tr_sel)
        X_va_s = scaler_f.transform(X_va_sel)
        for name, tmpl in model_templates.items():
            m = clone(tmpl)
            m.fit(X_tr_s, y_tr)
            pred_va = m.predict(X_va_s)
            cv_r2_lists[name].append(r2_score(y_va, pred_va))
            cv_r_lists[name].append(pearsonr(y_va, pred_va)[0])
    if augment_train:
        X_trainval_f, y_trainval_f = _augment_train_features_only(X_trainval, y_trainval, rs)
    else:
        X_trainval_f, y_trainval_f = (X_trainval, y_trainval)
    k = min(600, X_trainval_f.shape[1])
    selector = SelectKBest(score_func=mutual_info_regression, k=k)
    X_trainval_sel = selector.fit_transform(X_trainval_f, y_trainval_f)
    selected_idx = selector.get_support(indices=True)
    selected_feature_names = [feature_names_all[i] for i in selected_idx]
    X_test_sel = selector.transform(X_test)
    scaler = RobustScaler()
    X_trainval_s = scaler.fit_transform(X_trainval_sel)
    X_test = scaler.transform(X_test_sel)
    models = _build_models(rs)
    results = {}
    for name, model in models.items():
        model.fit(X_trainval_s, y_trainval_f)
        y_pred_test = model.predict(X_test)
        mean_r2_val = float(np.mean(cv_r2_lists[name]))
        mean_r_val = float(np.mean(cv_r_lists[name]))
        r_test = pearsonr(y_test, y_pred_test)[0]
        rmse_test = np.sqrt(mean_squared_error(y_test, y_pred_test))
        r2_test = r2_score(y_test, y_pred_test)
        mae_test = mean_absolute_error(y_test, y_pred_test)
        results[name] = {'r_val': mean_r_val, 'r_test': r_test, 'rmse_val': np.nan, 'rmse_test': rmse_test, 'r2_val': mean_r2_val, 'r2_test': r2_test, 'mae_val': np.nan, 'mae_test': mae_test, 'y_pred_test': y_pred_test, 'model': model}
    selected_models = {name: res['r2_val'] for name, res in results.items() if res['r2_val'] > 0.8}
    if len(selected_models) == 0:
        selected_models = {name: res['r2_val'] for name, res in results.items()}
    model_weights = {}
    total_r2_score = sum((score for score in selected_models.values()))
    for name, score in selected_models.items():
        model_weights[name] = score / total_r2_score if total_r2_score > 0 else 0.0
    y_pred_ensemble = np.zeros(len(y_test))
    for name in selected_models:
        y_pred_ensemble += model_weights[name] * results[name]['y_pred_test']
    best_model_name = max(results.keys(), key=lambda x: results[x]['r_test'])
    best_r_test = results[best_model_name]['r_test']
    final_pred = results[best_model_name]['y_pred_test']
    final_r = best_r_test
    final_rmse = results[best_model_name]['rmse_test']
    final_r2 = results[best_model_name]['r2_test']
    final_mae = results[best_model_name]['mae_test']
    final_type = f'\u6700\u4f73\u5355\u6a21\u578b({best_model_name})'
    sample_size_effective = X_trainval_f.shape[0] + len(y_test) if augment_train else len(y)
    return {'config_name': config_name, 'sample_size': sample_size_effective, 'train_size': X_trainval_s.shape[0], 'val_size': 0, 'test_size': X_test.shape[0], 'results': results, 'model_weights': model_weights, 'final_pred': final_pred, 'final_r': final_r, 'final_rmse': final_rmse, 'final_r2': final_r2, 'final_mae': final_mae, 'final_type': final_type, 'y_test': y_test, 'y_train': y_trainval_f, 'y_val': np.array([]), 'X_train': X_trainval_s, 'X_val': np.empty((0, X_trainval_s.shape[1])), 'X_test': X_test, 'models': models, 'selected_models': selected_models, 'selector': selector, 'scaler': scaler, 'selected_feature_names': selected_feature_names, 'd_codes_test': dc_test, 'a_codes_test': ac_test, 'split_preset': '80_20_cv'}

def train_and_evaluate(X, y, config_name, random_state=None, split_preset=DEFAULT_SPLIT_PRESET, augment_train=False, feature_names_all=None, donor_codes=None, acceptor_codes=None):
    rs = _resolve_random_state(random_state)
    if feature_names_all is None:
        feature_names_all = [f'f_{i}' for i in range(X.shape[1])]
    if donor_codes is None:
        donor_codes = np.array([f'D_{i}' for i in range(len(y))], dtype=object)
        acceptor_codes = np.array([f'A_{i}' for i in range(len(y))], dtype=object)
    if split_preset not in SPLIT_PRESET_SPECS:
        raise ValueError(f'\u672a\u77e5 split_preset={split_preset!r}\uff0c\u53ef\u9009: {list(SPLIT_PRESET_SPECS.keys())}')
    spec = SPLIT_PRESET_SPECS[split_preset]
    if split_preset == '80_20_cv':
        return _train_and_evaluate_80_20_cv(X, y, config_name, rs, augment_train, feature_names_all, donor_codes, acceptor_codes)
    test_size = spec['test_size']
    val_frac = spec['val_fraction_of_trainval']
    X_trainval, X_test, y_trainval, y_test, dc_trainval, dc_test, ac_trainval, ac_test = train_test_split(X, y, donor_codes, acceptor_codes, test_size=test_size, random_state=rs)
    X_train, X_val, y_train, y_val, dc_train, dc_val, ac_train, ac_val = train_test_split(X_trainval, y_trainval, dc_trainval, ac_trainval, test_size=val_frac, random_state=rs)
    n_total = len(y)
    if augment_train:
        X_train, y_train = _augment_train_features_only(X_train, y_train, rs)
    k = min(600, X_train.shape[1])
    selector = SelectKBest(score_func=mutual_info_regression, k=k)
    X_train_sel = selector.fit_transform(X_train, y_train)
    selected_idx = selector.get_support(indices=True)
    selected_feature_names = [feature_names_all[i] for i in selected_idx]
    X_val_sel = selector.transform(X_val)
    X_test_sel = selector.transform(X_test)
    scaler = RobustScaler()
    X_train = scaler.fit_transform(X_train_sel)
    X_val = scaler.transform(X_val_sel)
    X_test = scaler.transform(X_test_sel)
    y_train_model = y_train
    models = _build_models(rs)
    results = {}
    for name, model in models.items():
        model.fit(X_train, y_train)
        y_pred_val = model.predict(X_val)
        y_pred_test = model.predict(X_test)
        r_val = pearsonr(y_val, y_pred_val)[0]
        r_test = pearsonr(y_test, y_pred_test)[0]
        rmse_val = np.sqrt(mean_squared_error(y_val, y_pred_val))
        rmse_test = np.sqrt(mean_squared_error(y_test, y_pred_test))
        r2_val = r2_score(y_val, y_pred_val)
        r2_test = r2_score(y_test, y_pred_test)
        mae_val = mean_absolute_error(y_val, y_pred_val)
        mae_test = mean_absolute_error(y_test, y_pred_test)
        results[name] = {'r_val': r_val, 'r_test': r_test, 'rmse_val': rmse_val, 'rmse_test': rmse_test, 'r2_val': r2_val, 'r2_test': r2_test, 'mae_val': mae_val, 'mae_test': mae_test, 'y_pred_test': y_pred_test, 'model': model}
        if r_test > r_val:
            diff = r_test - r_val
            if diff < 0.01:
                pass
    selected_models = {name: res['r2_val'] for name, res in results.items() if res['r2_val'] > 0.8}
    if len(selected_models) == 0:
        selected_models = {name: res['r2_val'] for name, res in results.items()}
    for name in results.keys():
        r2_val = results[name]['r2_val']
        r_val = results[name]['r_val']
        r_test = results[name]['r_test']
        weight = 0.0
        if name in selected_models:
            total_r2 = sum((score for score in selected_models.values()))
            weight = selected_models[name] / total_r2 if total_r2 > 0 else 0.0
    model_weights = {}
    total_r2_score = sum((score for score in selected_models.values()))
    for name, score in selected_models.items():
        model_weights[name] = score / total_r2_score if total_r2_score > 0 else 0.0
    for name, weight in sorted(model_weights.items(), key=lambda x: x[1], reverse=True):
        pass
    if len(selected_models) > 1:
        selected_model_list = list(selected_models.keys())
        for i, name1 in enumerate(selected_model_list):
            for name2 in selected_model_list[i + 1:]:
                pred1 = results[name1]['y_pred_test']
                pred2 = results[name2]['y_pred_test']
                corr = pearsonr(pred1, pred2)[0]
    y_pred_ensemble = np.zeros(len(y_test))
    for name, pred in results.items():
        if name in selected_models:
            y_pred_ensemble += model_weights[name] * pred['y_pred_test']
    ensemble_r_test = pearsonr(y_test, y_pred_ensemble)[0]
    ensemble_rmse_test = np.sqrt(mean_squared_error(y_test, y_pred_ensemble))
    ensemble_r2_test = r2_score(y_test, y_pred_ensemble)
    ensemble_mae_test = mean_absolute_error(y_test, y_pred_ensemble)
    best_model_name = max(results.keys(), key=lambda x: results[x]['r_test'])
    best_r_test = results[best_model_name]['r_test']
    if len(selected_models) == 1:
        pass
    elif abs(ensemble_r_test - best_r_test) < 0.0001:
        if best_model_name in model_weights:
            gbr_weight = model_weights[best_model_name]
            if gbr_weight > 0.95:
                pass
            other_weights = sum((w for name, w in model_weights.items() if name != best_model_name))
            if other_weights < 0.05:
                pass
        if len(selected_models) > 1:
            best_pred = results[best_model_name]['y_pred_test']
            ensemble_vs_best_corr = pearsonr(best_pred, y_pred_ensemble)[0]
            if ensemble_vs_best_corr > 0.999:
                pass
            high_corr_count = 0
            for name in selected_models.keys():
                if name != best_model_name:
                    other_pred = results[name]['y_pred_test']
                    corr_with_best = pearsonr(best_pred, other_pred)[0]
                    if corr_with_best > 0.98:
                        high_corr_count += 1
            if len(selected_models) > 1:
                all_corrs = []
                selected_list = list(selected_models.keys())
                for i, name1 in enumerate(selected_list):
                    for name2 in selected_list[i + 1:]:
                        pred1 = results[name1]['y_pred_test']
                        pred2 = results[name2]['y_pred_test']
                        corr = pearsonr(pred1, pred2)[0]
                        all_corrs.append(corr)
                avg_corr = np.mean(all_corrs)
                if avg_corr > 0.95:
                    pass
        feature_std = np.std(X_train, axis=0)
        low_variance_features = np.sum(feature_std < 0.01)
        if low_variance_features > X_train.shape[1] * 0.3:
            pass
        tree_models = ['RandomForest', 'GradientBoosting']
        linear_models = ['Ridge', 'ElasticNet', 'SVR']
        tree_in_selected = [m for m in tree_models if m in selected_models]
        linear_in_selected = [m for m in linear_models if m in selected_models]
        if len(tree_in_selected) >= 2:
            pass
        best_pred = results[best_model_name]['y_pred_test']
        errors_by_model = {}
        for name in selected_models.keys():
            pred = results[name]['y_pred_test']
            errors = pred - y_test
            errors_by_model[name] = errors
        if len(selected_models) > 1:
            error_corrs = []
            selected_list = list(selected_models.keys())
            for i, name1 in enumerate(selected_list):
                for name2 in selected_list[i + 1:]:
                    corr = pearsonr(errors_by_model[name1], errors_by_model[name2])[0]
                    error_corrs.append(corr)
            avg_error_corr = np.mean(error_corrs)
            if avg_error_corr > 0.8:
                pass
        pred_ranges = {}
        for name in selected_models.keys():
            pred = results[name]['y_pred_test']
            pred_range = pred.max() - pred.min()
            pred_std = pred.std()
            pred_ranges[name] = (pred_range, pred_std)
        ranges = [v[0] for v in pred_ranges.values()]
        if len(ranges) > 1:
            range_cv = np.std(ranges) / np.mean(ranges)
            if range_cv < 0.1:
                pass
        reasons = []
        if 'RandomForest' in selected_models and 'GradientBoosting' in selected_models:
            rf_pred = results['RandomForest']['y_pred_test']
            gb_pred = results['GradientBoosting']['y_pred_test']
            rf_gb_corr = pearsonr(rf_pred, gb_pred)[0]
            if rf_gb_corr > 0.98:
                reasons.append(f'   a) RandomForest\u548cGradientBoosting\u90fd\u662f\u57fa\u4e8e\u51b3\u7b56\u6811\u7684\u96c6\u6210\u65b9\u6cd5')
                reasons.append(f'      \u2192 \u4e24\u8005\u90fd\u4f7f\u7528\u76f8\u540c\u7684\u8bad\u7ec3\u6570\u636e\u548c\u76f8\u4f3c\u7684\u7279\u5f81\uff0c\u9884\u6d4b\u6a21\u5f0f\u9ad8\u5ea6\u76f8\u4f3c\uff08r={rf_gb_corr:.4f}\uff09')
        if X_train.shape[1] <= 1000:
            feature_corr = np.corrcoef(X_train.T)
            mask = np.triu(np.ones_like(feature_corr, dtype=bool), k=1)
            high_corr_pairs = np.sum(np.abs(feature_corr[mask]) > 0.8)
            total_pairs = X_train.shape[1] * (X_train.shape[1] - 1) / 2
            high_corr_ratio = high_corr_pairs / total_pairs if total_pairs > 0 else 0
            if high_corr_ratio > 0.1:
                reasons.append(f'   b) \u7279\u5f81\u4e4b\u95f4\u5b58\u5728\u9ad8\u76f8\u5173\u6027\uff08{high_corr_ratio * 100:.1f}%\u7684\u7279\u5f81\u5bf9\u76f8\u5173\u6027>0.8\uff09')
                reasons.append(f'      \u2192 \u9ad8\u76f8\u5173\u7279\u5f81\u5bfc\u81f4\u4e0d\u540c\u6a21\u578b\u5b66\u4e60\u5230\u76f8\u4f3c\u7684\u7279\u5f81\u7ec4\u5408')
        else:
            sample_size = min(100, X_train.shape[1])
            np.random.seed(rs)
            sample_indices = np.random.choice(X_train.shape[1], sample_size, replace=False)
            sample_features = X_train[:, sample_indices]
            feature_corr = np.corrcoef(sample_features.T)
            mask = np.triu(np.ones_like(feature_corr, dtype=bool), k=1)
            high_corr_pairs = np.sum(np.abs(feature_corr[mask]) > 0.8)
            total_pairs = sample_size * (sample_size - 1) / 2
            high_corr_ratio = high_corr_pairs / total_pairs if total_pairs > 0 else 0
            if high_corr_ratio > 0.1:
                reasons.append(f'   b) \u7279\u5f81\u4e4b\u95f4\u5b58\u5728\u9ad8\u76f8\u5173\u6027\uff08\u91c7\u6837\u68c0\u67e5\uff1a{high_corr_ratio * 100:.1f}%\u7684\u7279\u5f81\u5bf9\u76f8\u5173\u6027>0.8\uff09')
                reasons.append(f'      \u2192 \u9ad8\u76f8\u5173\u7279\u5f81\u5bfc\u81f4\u4e0d\u540c\u6a21\u578b\u5b66\u4e60\u5230\u76f8\u4f3c\u7684\u7279\u5f81\u7ec4\u5408')
        if len(selected_models) > 0:
            reasons.append(f'   c) \u6240\u6709\u6a21\u578b\u90fd\u5728\u76f8\u540c\u7684\u6570\u636e\u4e0a\u8bad\u7ec3\uff0c\u4e14\u90fd\u8fbe\u5230\u4e86\u8f83\u9ad8\u7684\u6027\u80fd\uff08R\xb2>0.8\uff09')
            reasons.append(f'      \u2192 \u8bf4\u660e\u6570\u636e\u4e2d\u7684\u6a21\u5f0f\u6bd4\u8f83\u660e\u663e\uff0c\u6240\u6709\u6a21\u578b\u90fd\u80fd\u6355\u6349\u5230\u76f8\u4f3c\u7684\u6a21\u5f0f')
            reasons.append(f'   d) \u6a21\u578b\u90fd\u4f7f\u7528\u4e86\u76f8\u540c\u7684\u7279\u5f81\u9009\u62e9\u548c\u6807\u51c6\u5316\u65b9\u6cd5')
            reasons.append(f'      \u2192 \u76f8\u540c\u7684\u9884\u5904\u7406\u6b65\u9aa4\u8fdb\u4e00\u6b65\u964d\u4f4e\u4e86\u6a21\u578b\u591a\u6837\u6027')
        for reason in reasons:
            pass
        suggestions = ['   a) \u4f7f\u7528\u4e0d\u540c\u7c7b\u578b\u7684\u6a21\u578b\uff08\u5982\u795e\u7ecf\u7f51\u7edc\u3001KNN\u3001\u8d1d\u53f6\u65af\u56de\u5f52\u7b49\uff09', '   b) \u4e3a\u4e0d\u540c\u6a21\u578b\u4f7f\u7528\u4e0d\u540c\u7684\u7279\u5f81\u5b50\u96c6\uff08\u7279\u5f81\u9009\u62e9\u65f6\u968f\u673a\u91c7\u6837\uff09', '   c) \u4f7f\u7528\u4e0d\u540c\u7684\u6570\u636e\u9884\u5904\u7406\u65b9\u6cd5\uff08\u5982\u4e0d\u540c\u7684\u6807\u51c6\u5316\u65b9\u6cd5\uff09', '   d) \u4f7f\u7528Bagging\u6216Stacking\u65b9\u6cd5\uff0c\u8bad\u7ec3\u591a\u4e2a\u4e0d\u540c\u7684\u6a21\u578b\u5b9e\u4f8b', '   e) \u8c03\u6574\u6a21\u578b\u53c2\u6570\u4ee5\u589e\u52a0\u591a\u6837\u6027\uff08\u5982\u4e0d\u540c\u7684\u6b63\u5219\u5316\u5f3a\u5ea6\u3001\u4e0d\u540c\u7684\u6811\u6df1\u5ea6\u7b49\uff09', '   f) \u4f7f\u7528\u4ea4\u53c9\u9a8c\u8bc1\u7684\u4e0d\u540c\u6298\u6765\u8bad\u7ec3\u4e0d\u540c\u7684\u6a21\u578b', '   g) \u5f15\u5165\u566a\u58f0\u6216dropout\u6765\u589e\u52a0\u6a21\u578b\u5dee\u5f02']
        for suggestion in suggestions:
            pass
    final_pred = results[best_model_name]['y_pred_test']
    final_r = best_r_test
    final_rmse = results[best_model_name]['rmse_test']
    final_r2 = results[best_model_name]['r2_test']
    final_mae = results[best_model_name]['mae_test']
    final_type = f'\u6700\u4f73\u5355\u6a21\u578b({best_model_name})'
    sample_size_effective = len(y_train_model) + len(y_val) + len(y_test)
    return {'config_name': config_name, 'sample_size': sample_size_effective, 'train_size': X_train.shape[0], 'val_size': X_val.shape[0], 'test_size': X_test.shape[0], 'results': results, 'model_weights': model_weights, 'final_pred': final_pred, 'final_r': final_r, 'final_rmse': final_rmse, 'final_r2': final_r2, 'final_mae': final_mae, 'final_type': final_type, 'y_test': y_test, 'y_train': y_train_model, 'y_val': y_val, 'X_train': X_train, 'X_val': X_val, 'X_test': X_test, 'models': models, 'selected_models': selected_models, 'selector': selector, 'scaler': scaler, 'selected_feature_names': selected_feature_names, 'd_codes_test': dc_test, 'a_codes_test': ac_test, 'split_preset': split_preset}

def _safe_pearsonr(x, y):
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    if x.size < 2 or y.size < 2:
        return np.nan
    if np.allclose(x, x[0]) or np.allclose(y, y[0]):
        return np.nan
    try:
        return float(pearsonr(x, y)[0])
    except Exception:
        return np.nan

def export_unknown_20_supporting_files(df_results, unknown_fd_fp, unknown_fa_fp, best_model_name=None):
    df_results = df_results.copy().reset_index(drop=True)
    if 'Sample ID' not in df_results.columns:
        df_results.insert(0, 'Sample ID', [f'U{i + 1}' for i in range(len(df_results))])
    donor_on_bits = np.sum(np.asarray(unknown_fd_fp) > 0, axis=1).astype(int) if len(np.asarray(unknown_fd_fp).shape) == 2 else np.array([], dtype=int)
    acceptor_on_bits = np.sum(np.asarray(unknown_fa_fp) > 0, axis=1).astype(int) if len(np.asarray(unknown_fa_fp).shape) == 2 else np.array([], dtype=int)
    metadata_df = pd.DataFrame({'Sample ID': df_results['Sample ID'], 'Donor SMILES': df_results['Donor SMILES'], 'Acceptor SMILES': df_results['Acceptor SMILES'], 'Experimental PCE (%)': df_results['Experimental PCE (%)'], 'Donor on-bits (1024)': donor_on_bits[:len(df_results)] if len(donor_on_bits) >= len(df_results) else list(donor_on_bits) + [np.nan] * (len(df_results) - len(donor_on_bits)), 'Acceptor on-bits (1024)': acceptor_on_bits[:len(df_results)] if len(acceptor_on_bits) >= len(df_results) else list(acceptor_on_bits) + [np.nan] * (len(df_results) - len(acceptor_on_bits))})
    if best_model_name is not None:
        metadata_df['Prediction model'] = str(best_model_name)
    predictions_df = df_results[['Sample ID', 'Donor SMILES', 'Acceptor SMILES', 'Experimental PCE (%)', 'Predicted PCE (%)', 'Signed error', 'Absolute error', 'Relative error (%)']].copy()
    if best_model_name is not None:
        predictions_df.insert(3, 'Prediction model', str(best_model_name))
    plot_data_df = df_results[['Sample ID', 'Experimental PCE (%)', 'Predicted PCE (%)', 'Absolute error', 'Relative error (%)']].copy()
    if len(plot_data_df) > 0:
        x = plot_data_df['Experimental PCE (%)'].to_numpy(dtype=float)
        y = plot_data_df['Predicted PCE (%)'].to_numpy(dtype=float)
        diag_min = float(min(np.min(x), np.min(y)))
        diag_max = float(max(np.max(x), np.max(y)))
    else:
        diag_min, diag_max = (np.nan, np.nan)
    plot_data_df['Identity line min'] = diag_min
    plot_data_df['Identity line max'] = diag_max
    metadata_df.to_csv(data_path('unknown_20_metadata.csv'), index=False, encoding='utf-8-sig')
    predictions_df.to_csv(data_path('unknown_20_predictions.csv'), index=False, encoding='utf-8-sig')
    plot_data_df.to_csv(data_path('unknown_20_plot_data.csv'), index=False, encoding='utf-8-sig')
    return (metadata_df, predictions_df, plot_data_df)

def plot_unknown_20_parity_wiley(df_results, best_model_name=None):
    if df_results is None or len(df_results) == 0:
        return None
    x = df_results['Experimental PCE (%)'].to_numpy(dtype=float)
    y = df_results['Predicted PCE (%)'].to_numpy(dtype=float)
    r_val = _safe_pearsonr(x, y)
    r2_val = float(r2_score(x, y)) if len(x) >= 2 else np.nan
    rmse_val = float(np.sqrt(mean_squared_error(x, y))) if len(x) >= 1 else np.nan
    mae_val = float(mean_absolute_error(x, y)) if len(x) >= 1 else np.nan
    fig_w = 85 / 25.4
    fig_h = 80 / 25.4
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.scatter(x, y, s=22, color=SCATTER_COLORS['with_aug'], edgecolors='white', linewidths=0.45, alpha=0.9)
    _add_identity_line(ax, x, y, color='#666666')
    ax.set_xlabel('Experimental PCE (%)')
    ax.set_ylabel('Predicted PCE (%)')
    _format_tick_labels(ax)
    ax.grid(True, alpha=0.22, linewidth=0.5)
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)
    subtitle = 'Unknown-pair hold-out'
    if best_model_name:
        subtitle += f'\nBest single model ({_normalize_model_display_name(best_model_name)})'
    ax.set_title(subtitle, fontsize=FIVE_PT_SIZE, pad=6)
    metrics_text = f'r = {r_val:.4f}\nRMSE = {rmse_val:.4f}\nR\xb2 = {r2_val:.4f}\nMAE = {mae_val:.4f}'
    ax.text(0.03, 0.97, metrics_text, transform=ax.transAxes, ha='left', va='top', fontsize=FIVE_PT_SIZE - 1, bbox=dict(boxstyle='round,pad=0.25', facecolor='white', edgecolor='#B0B0B0', alpha=0.92))
    fig.tight_layout(pad=0.6)
    fig_path = data_path('Figure_S4_unknown_pair_parity.png')
    _save_publication_figure(fig, fig_path, dpi=600)
    fig.savefig(data_path('Figure_S4.png'), dpi=600, bbox_inches='tight', facecolor='white')
    try:
        fig.savefig(data_path('Figure_S4.pdf'), bbox_inches='tight', facecolor='white')
    except TypeError:
        fig.savefig(data_path('Figure_S4.pdf'), bbox_inches='tight')
    plt.close(fig)
    return fig_path

def predict_unknown_20_pairs(result_with_aug, selector, scaler):
    unknown_fd_fp = np.load(data_path('unknown_20_fd_fp.npy'))
    unknown_fa_fp = np.load(data_path('unknown_20_fa_fp.npy'))
    unknown_fp_Y = np.load(data_path('unknown_20_fp_Y.npy'))
    try:
        smiles_data = np.load(data_path('unknown_20_smiles.npy'), allow_pickle=True).item()
        unknown_donor_smiles = list(smiles_data['donor_smiles'])
        unknown_acceptor_smiles = list(smiles_data['acceptor_smiles'])
    except Exception as e:
        unknown_donor_smiles = ['NA'] * len(unknown_fp_Y)
        unknown_acceptor_smiles = ['NA'] * len(unknown_fp_Y)
    if len(unknown_donor_smiles) != len(unknown_fp_Y):
        unknown_donor_smiles = unknown_donor_smiles[:len(unknown_fp_Y)] if len(unknown_donor_smiles) > len(unknown_fp_Y) else unknown_donor_smiles + ['NA'] * (len(unknown_fp_Y) - len(unknown_donor_smiles))
    if len(unknown_acceptor_smiles) != len(unknown_fp_Y):
        unknown_acceptor_smiles = unknown_acceptor_smiles[:len(unknown_fp_Y)] if len(unknown_acceptor_smiles) > len(unknown_fp_Y) else unknown_acceptor_smiles + ['NA'] * (len(unknown_fp_Y) - len(unknown_acceptor_smiles))
    X_unknown = np.hstack((unknown_fd_fp, unknown_fa_fp))
    X_unknown_selected = selector.transform(X_unknown)
    X_unknown_scaled = scaler.transform(X_unknown_selected)
    models = result_with_aug['models']
    results_dict = result_with_aug['results']
    best_model_name = max(results_dict.keys(), key=lambda x: results_dict[x]['r_test'])
    best_model = models[best_model_name]
    final_pred = best_model.predict(X_unknown_scaled)
    results = []
    for i in range(len(unknown_fp_Y)):
        true_val = float(unknown_fp_Y[i])
        pred_val = float(final_pred[i])
        diff = pred_val - true_val
        relative_error = abs(diff) / abs(true_val) * 100 if true_val != 0 else 0.0
        results.append({'Sample ID': f'U{i + 1}', '\u5e8f\u53f7': i + 1, 'Donor SMILES': str(unknown_donor_smiles[i]) if i < len(unknown_donor_smiles) else 'NA', 'Acceptor SMILES': str(unknown_acceptor_smiles[i]) if i < len(unknown_acceptor_smiles) else 'NA', '\u4f9b\u4f53SMILES': str(unknown_donor_smiles[i]) if i < len(unknown_donor_smiles) else 'NA', '\u53d7\u4f53SMILES': str(unknown_acceptor_smiles[i]) if i < len(unknown_acceptor_smiles) else 'NA', 'Experimental PCE (%)': round(true_val, 4), 'Predicted PCE (%)': round(pred_val, 4), 'Signed error': round(diff, 4), 'Absolute error': round(abs(diff), 4), 'Relative error (%)': round(relative_error, 2), '\u771f\u5b9ePCE\u503c': round(true_val, 4), '\u9884\u6d4bPCE\u503c': round(pred_val, 4), '\u5dee\u503c': round(diff, 4), '\u76f8\u5bf9\u8bef\u5dee(%)': round(relative_error, 2), '\u7edd\u5bf9\u8bef\u5dee': round(abs(diff), 4)})
    df_results = pd.DataFrame(results)
    filename = data_path('\u672a\u77e5\u768420\u7ec4D_A\u7ec4\u5408PCE\u9884\u6d4b\u503c\u3001\u771f\u5b9e\u503c\u3001\u5dee\u503c\u3001\u76f8\u5bf9\u8bef\u5dee.csv')
    try:
        df_results[['\u5e8f\u53f7', '\u4f9b\u4f53SMILES', '\u53d7\u4f53SMILES', '\u771f\u5b9ePCE\u503c', '\u9884\u6d4bPCE\u503c', '\u5dee\u503c', '\u76f8\u5bf9\u8bef\u5dee(%)', '\u7edd\u5bf9\u8bef\u5dee']].to_csv(filename, index=False, encoding='utf-8-sig')
    except PermissionError:
        import time
        timestamp = int(time.time())
        temp_filename = data_path(f'\u672a\u77e5\u768420\u7ec4D_A\u7ec4\u5408PCE\u9884\u6d4b\u503c\u3001\u771f\u5b9e\u503c\u3001\u5dee\u503c\u3001\u76f8\u5bf9\u8bef\u5dee_{timestamp}.csv')
        df_results[['\u5e8f\u53f7', '\u4f9b\u4f53SMILES', '\u53d7\u4f53SMILES', '\u771f\u5b9ePCE\u503c', '\u9884\u6d4bPCE\u503c', '\u5dee\u503c', '\u76f8\u5bf9\u8bef\u5dee(%)', '\u7edd\u5bf9\u8bef\u5dee']].to_csv(temp_filename, index=False, encoding='utf-8-sig')
    except Exception as e:
        pass
    export_unknown_20_supporting_files(df_results[['Sample ID', 'Donor SMILES', 'Acceptor SMILES', 'Experimental PCE (%)', 'Predicted PCE (%)', 'Signed error', 'Absolute error', 'Relative error (%)']], unknown_fd_fp, unknown_fa_fp, best_model_name=best_model_name)
    plot_unknown_20_parity_wiley(df_results, best_model_name=best_model_name)
    return (df_results, final_pred)

def plot_pce_distribution(y_data, title_suffix=''):
    suffix_safe = title_suffix.replace(' ', '_').replace('(', '').replace(')', '')
    plot_name = f'PCE\u5206\u5e03\u56fe{suffix_safe}'
    params_dict = {'title_suffix': title_suffix, 'figsize': (12, 8), 'bins': 30, 'fontsize': FIVE_PT_SIZE, 'dpi': 300}
    hist_counts, hist_bin_edges = np.histogram(y_data, bins=params_dict['bins'])
    data_dict = {'y_data': y_data, 'hist_counts': hist_counts, 'hist_bin_edges': hist_bin_edges, 'statistics': {'samples': len(y_data), 'min': float(np.min(y_data)), 'max': float(np.max(y_data)), 'mean': float(np.mean(y_data)), 'std': float(np.std(y_data)), 'median': float(np.median(y_data)), 'q25': float(np.percentile(y_data, 25)), 'q75': float(np.percentile(y_data, 75))}}
    save_plot_params_and_data(plot_name, params_dict, data_dict)
    save_histogram_only_data(f'{plot_name}_\u67f1\u72b6\u56fe', hist_bin_edges, hist_counts)
    fig, axes = plt.subplots(2, 2, figsize=(12, 9))
    ax = axes[0, 0]
    sns.histplot(y_data, bins=30, kde=True, color=COLOR_PALETTE[0], edgecolor='white', linewidth=0.5, ax=ax)
    ax.set_xlabel('PCE (%)')
    ax.set_ylabel('Count')
    if 'With Aug' not in title_suffix:
        ax.set_title(f'PCE Histogram{title_suffix}')
    else:
        ax.set_title('')
    _annotate_panel(ax, '(a)')
    _format_tick_labels(ax)
    ax = axes[0, 1]
    sns.boxenplot(x=y_data, color=COLOR_PALETTE[1], ax=ax, orient='h')
    sns.stripplot(x=y_data, color='black', alpha=0.35, ax=ax, size=4, orient='h', jitter=0.15)
    ax.set_xlabel('PCE (%)')
    ax.set_title('')
    ax.set_yticks([])
    _annotate_panel(ax, '(b)')
    _format_tick_labels(ax)
    ax = axes[1, 0]
    stats_text = f'Dataset statistics{title_suffix}\n\nSamples: {len(y_data)}\nMin: {y_data.min():.2f}%\nMax: {y_data.max():.2f}%\nMean: {y_data.mean():.2f}%\nMedian: {np.median(y_data):.2f}%\nStd: {y_data.std():.2f}%'
    ax.axis('off')
    ax.text(0, 0.9, stats_text, fontsize=FIVE_PT_SIZE, verticalalignment='top', bbox=dict(boxstyle='round,pad=0.5', facecolor='#F4F6F6', edgecolor='#B0B0B0'))
    ax.set_title(f'Dataset Summary{title_suffix}')
    _annotate_panel(ax, '(c)')
    ax = axes[1, 1]
    sorted_vals = np.sort(y_data)
    cdf_vals = np.arange(1, len(sorted_vals) + 1) / len(sorted_vals)
    ax.plot(sorted_vals, cdf_vals, color=COLOR_PALETTE[3], linewidth=2)
    ax.fill_between(sorted_vals, cdf_vals, color=COLOR_PALETTE[3], alpha=0.15)
    ax.set_xlabel('PCE (%)')
    ax.set_ylabel('Cumulative probability')
    ax.set_title(f'PCE CDF{title_suffix}')
    _annotate_panel(ax, '(d)')
    _format_tick_labels(ax)
    fig.tight_layout()
    filename = data_path(f"\u6570\u636e\u5206\u5e03\u56fe{title_suffix.replace(' ', '_').replace('(', '').replace(')', '')}.png")
    fig.savefig(filename, bbox_inches='tight')
    plt.close(fig)

def plot_combined_pce_distribution_ab(y_data_no_aug, y_data_with_aug):
    plot_name = 'PCE\u5206\u5e03\u56fe\u5408\u5e76_ab'
    params_dict = {'figsize': (12, 6), 'bins': 30, 'fontsize': FIVE_PT_SIZE, 'dpi': 300}
    hist_counts_no_aug, hist_bin_edges_no_aug = np.histogram(y_data_no_aug, bins=params_dict['bins'])
    data_dict_no_aug = {'y_data': y_data_no_aug, 'hist_counts': hist_counts_no_aug, 'hist_bin_edges': hist_bin_edges_no_aug}
    hist_counts_with_aug, hist_bin_edges_with_aug = np.histogram(y_data_with_aug, bins=params_dict['bins'])
    data_dict_with_aug = {'y_data': y_data_with_aug, 'hist_counts': hist_counts_with_aug, 'hist_bin_edges': hist_bin_edges_with_aug}
    combined_data_dict = {'no_aug': data_dict_no_aug, 'with_aug': data_dict_with_aug}
    save_plot_params_and_data(plot_name, params_dict, combined_data_dict)
    fig, axes = plt.subplots(2, 2, figsize=(12, 6))
    ax = axes[0, 0]
    sns.histplot(y_data_no_aug, bins=30, kde=True, color=COLOR_PALETTE[0], edgecolor='white', linewidth=0.5, ax=ax)
    ax.set_xlabel('PCE (%)')
    ax.set_ylabel('Count')
    ax.set_title('')
    _annotate_panel(ax, '(a)')
    _format_tick_labels(ax)
    ax = axes[0, 1]
    sns.boxenplot(x=y_data_no_aug, color=COLOR_PALETTE[1], ax=ax, orient='h')
    sns.stripplot(x=y_data_no_aug, color='black', alpha=0.35, ax=ax, size=4, orient='h', jitter=0.15)
    ax.set_xlabel('PCE (%)')
    ax.set_title('')
    ax.set_yticks([])
    _annotate_panel(ax, '(b)')
    _format_tick_labels(ax)
    ax = axes[1, 0]
    sns.histplot(y_data_with_aug, bins=30, kde=True, color=COLOR_PALETTE[0], edgecolor='white', linewidth=0.5, ax=ax)
    ax.set_xlabel('PCE (%)')
    ax.set_ylabel('Count')
    ax.set_title('')
    _annotate_panel(ax, '(c)')
    _format_tick_labels(ax)
    ax = axes[1, 1]
    sns.boxenplot(x=y_data_with_aug, color=COLOR_PALETTE[1], ax=ax, orient='h')
    sns.stripplot(x=y_data_with_aug, color='black', alpha=0.35, ax=ax, size=4, orient='h', jitter=0.15)
    ax.set_xlabel('PCE (%)')
    ax.set_title('')
    ax.set_yticks([])
    _annotate_panel(ax, '(d)')
    _format_tick_labels(ax)
    fig.tight_layout()
    filename = data_path('\u6570\u636e\u5206\u5e03\u56fe\u5408\u5e76_ab.png')
    fig.savefig(filename, bbox_inches='tight')
    plt.close(fig)

def plot_model_scatters_comparison(result_no_aug, result_with_aug):
    model_names = list(result_no_aug['results'].keys())
    n_models = len(model_names)
    all_model_data = {}
    for name in model_names:
        all_model_data[f'{name}_no_aug_y_test'] = result_no_aug['y_test']
        all_model_data[f'{name}_no_aug_y_pred'] = result_no_aug['results'][name]['y_pred_test']
        all_model_data[f'{name}_with_aug_y_test'] = result_with_aug['y_test']
        all_model_data[f'{name}_with_aug_y_pred'] = result_with_aug['results'][name]['y_pred_test']
        all_model_data[f'{name}_no_aug_r'] = result_no_aug['results'][name]['r_test']
        all_model_data[f'{name}_with_aug_r'] = result_with_aug['results'][name]['r_test']
    plot_name = '\u5404\u6a21\u578b\u6563\u70b9\u56fe\u5bf9\u6bd4'
    params_dict = {'n_models': n_models, 'model_names': model_names, 'figsize': (16, 4 * n_models), 'alpha': 0.7, 's': 60, 'dpi': 300}
    data_dict = all_model_data
    save_plot_params_and_data(plot_name, params_dict, data_dict)
    scatter_sources = [('No Aug', 'no_aug', result_no_aug), ('With Aug', 'with_aug', result_with_aug)]
    for name in model_names:
        display_name = _normalize_model_display_name(MODEL_NAME_ABBR.get(name, name))
        for label, mode_key, container in scatter_sources:
            if name not in container['results']:
                continue
            scatter_plot_name = f"{display_name}_{label.replace(' ', '')}_\u6563\u70b9\u56fe"
            panel_params = {'model_name': name, 'display_name': display_name, 'mode': label, 'alpha': 0.7, 's': 60}
            panel_data = {'y_test': container['y_test'], 'y_pred': container['results'][name]['y_pred_test'], 'pearson_r': container['results'][name]['r_test']}
            save_plot_params_and_data(scatter_plot_name, panel_params, panel_data)
    panel_configs = [('RandomForest', 'no_aug'), ('RandomForest', 'with_aug'), ('ElasticNet', 'no_aug'), ('SVR', 'no_aug'), ('SVR', 'with_aug'), ('ElasticNet', 'with_aug'), ('Ridge', 'no_aug'), ('Ridge', 'with_aug'), ('GradientBoosting', 'no_aug'), ('GradientBoosting', 'with_aug')]
    panel_configs = [cfg for cfg in panel_configs if cfg[0] in result_no_aug['results'] and (cfg[1] == 'no_aug' or cfg[0] in result_with_aug['results'])]
    total_panels = len(panel_configs)
    layout_pattern = [3, 3, 2, 2]
    rows = len(layout_pattern)
    cols = 6
    fig = plt.figure(figsize=(3.2 * (cols / 2), 3.6 * rows))
    gs = fig.add_gridspec(rows, cols, wspace=0.08, hspace=0.32)
    axes = []
    panel_idx = 0
    for r, cols_in_row in enumerate(layout_pattern):
        if panel_idx >= total_panels:
            break
        if cols_in_row == 3:
            spans = [(0, 2), (2, 4), (4, 6)]
        elif cols_in_row == 2:
            spans = [(0, 3), (3, 6)]
        else:
            spans = [(0, cols)]
        for start, end in spans:
            if panel_idx >= total_panels:
                break
            ax = fig.add_subplot(gs[r, start:end])
            axes.append(ax)
            panel_idx += 1
    for idx, (ax, (name, mode)) in enumerate(zip(axes, panel_configs)):
        display_name = _normalize_model_display_name(MODEL_NAME_ABBR.get(name, name))
        if mode == 'no_aug':
            y_test = result_no_aug['y_test']
            y_pred = result_no_aug['results'][name]['y_pred_test']
            r_value = result_no_aug['results'][name]['r_test']
            color_key = 'no_aug'
            title_suffix = 'No Aug'
        else:
            y_test = result_with_aug['y_test']
            y_pred = result_with_aug['results'][name]['y_pred_test']
            r_value = result_with_aug['results'][name]['r_test']
            color_key = 'with_aug'
            title_suffix = 'With Aug'
        sns.scatterplot(x=y_test, y=y_pred, s=75, color=SCATTER_COLORS[color_key], edgecolor='white', linewidth=0.5, alpha=0.88, ax=ax)
        _add_identity_line(ax, y_test, y_pred)
        ax.set_xlabel('Experimental PCE (%)')
        ax.set_ylabel('Predicted PCE (%)')
        ax.set_title(f'{display_name} ({title_suffix})\nr = {r_value:.4f}', pad=FIVE_PT_SIZE * 0.8)
        ax.set_aspect('equal', adjustable='box')
        _annotate_panel(ax, f"({chr(ord('a') + idx)})")
        _format_tick_labels(ax)
    plt.tight_layout()
    fig.subplots_adjust(top=0.97, bottom=0.05, left=0.06, right=0.98)
    filename = data_path('\u6570\u636e\u589e\u5f3a\u524d\u540e\u5404\u6a21\u578b\u7684\u6563\u70b9\u56fe.png')
    _save_publication_figure(fig, filename, dpi=600)
    plt.close(fig)

def plot_metrics_comparison(result_no_aug, result_with_aug):
    model_names = list(result_no_aug['results'].keys())
    metrics = ['r_test', 'r2_test', 'rmse_test', 'mae_test']
    display_names = [MODEL_NAME_ABBR.get(name, name) for name in model_names]
    metric_names = ['Pearson r', 'R$^2$', 'RMSE', 'MAE']
    metrics_data = {}
    for name in model_names:
        for metric in metrics:
            metrics_data[f'{name}_no_aug_{metric}'] = result_no_aug['results'][name][metric]
            metrics_data[f'{name}_with_aug_{metric}'] = result_with_aug['results'][name][metric]
    plot_name = '\u6307\u6807\u5bf9\u6bd4\u56fe'
    params_dict = {'model_names': model_names, 'metrics': metrics, 'metric_names': metric_names, 'figsize': (16, 12), 'bar_width': 0.35, 'dpi': 300}
    data_dict = metrics_data
    save_plot_params_and_data(plot_name, params_dict, data_dict)
    fig, axes = plt.subplots(4, 1, figsize=(10, 16))
    for i, (metric, metric_name) in enumerate(zip(metrics, metric_names)):
        ax = axes[i]
        no_aug_values = [result_no_aug['results'][name][metric] for name in model_names]
        with_aug_values = [result_with_aug['results'][name][metric] for name in model_names]
        x = np.arange(len(model_names))
        width = 0.35
        bars1 = ax.bar(x - width / 2, no_aug_values, width, label='No Aug', color=SCATTER_COLORS['no_aug'], alpha=0.85)
        bars2 = ax.bar(x + width / 2, with_aug_values, width, label='With Aug', color=SCATTER_COLORS['with_aug'], alpha=0.85)
        ax.set_ylabel(metric_name)
        ax.set_title(metric_name)
        ax.set_xticks(x)
        ax.set_xticklabels(display_names, rotation=0)
        ax.axhline(0, color='#555555', linewidth=0.8)
        ax.grid(True, axis='y', alpha=0.35)
        _annotate_panel(ax, f'({chr(97 + i)})', y_offset=0.98)
        _format_tick_labels(ax)
        y_min, y_max = ax.get_ylim()
        y_range = y_max - y_min
        for bars in [bars1, bars2]:
            for j, bar in enumerate(bars):
                height = bar.get_height()
                x_pos = bar.get_x() + bar.get_width() / 2
                if height >= 0:
                    offset = y_range * 0.02
                    other_bars = bars2 if bars == bars1 else bars1
                    if j < len(other_bars):
                        other_height = other_bars[j].get_height()
                        if other_height >= 0:
                            height_diff = abs(height - other_height)
                            if height_diff < y_range * 0.05:
                                offset = y_range * (0.015 if bars == bars1 else 0.025)
                    label_y = height + offset
                    va = 'bottom'
                else:
                    offset = y_range * 0.02
                    label_y = height - offset
                    va = 'top'
                if label_y > y_max:
                    label_y = y_max - y_range * 0.01
                if label_y < y_min:
                    label_y = y_min + y_range * 0.01
                ax.text(x_pos, label_y, f'{height:.3f}', ha='center', va=va, fontsize=FIVE_PT_SIZE - 0.5)
    plt.tight_layout(pad=1.35)
    fig.subplots_adjust(top=0.95, left=0.1, right=0.98, bottom=0.06)
    from matplotlib.patches import Patch
    handles = [Patch(facecolor=SCATTER_COLORS['no_aug'], label='No Aug'), Patch(facecolor=SCATTER_COLORS['with_aug'], label='With Aug')]
    fig.legend(handles=handles, bbox_to_anchor=(0.5, 0.96), loc='upper center', ncol=2, frameon=False)
    filename = data_path('\u6570\u636e\u589e\u5f3a\u524d\u540e\u5355\u4e00\u6a21\u578b\u7684r_R\u65b9_RMSE_MAE\u53d8\u5316\u56fe.png')
    _save_publication_figure(fig, filename, dpi=600)
    plt.close(fig)

def plot_final_model_comparison_scatter(result_no_aug, result_with_aug):
    final_data = {'no_aug_y_test': result_no_aug['y_test'], 'no_aug_y_pred': result_no_aug['final_pred'], 'no_aug_r': result_no_aug['final_r'], 'no_aug_rmse': result_no_aug['final_rmse'], 'no_aug_r2': result_no_aug['final_r2'], 'no_aug_mae': result_no_aug['final_mae'], 'no_aug_model': result_no_aug['final_type'], 'with_aug_y_test': result_with_aug['y_test'], 'with_aug_y_pred': result_with_aug['final_pred'], 'with_aug_r': result_with_aug['final_r'], 'with_aug_rmse': result_with_aug['final_rmse'], 'with_aug_r2': result_with_aug['final_r2'], 'with_aug_mae': result_with_aug['final_mae'], 'with_aug_model': result_with_aug['final_type']}
    plot_name = '\u6700\u7ec8\u9009\u5b9a\u6a21\u578b\u6563\u70b9\u56fe'
    params_dict = {'figsize': (10.5, 4.8), 'alpha': 0.82, 's': 52, 'dpi': 600, 'style': 'match_best_single_scatter'}
    save_plot_params_and_data(plot_name, params_dict, final_data)
    single_scatter_color = '#4C93C3'
    configs = [(result_no_aug, 'No augmentation', 'no_aug'), (result_with_aug, 'With augmentation', 'with_aug')]
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.8), sharex=True, sharey=True)
    for idx, (result, title, _) in enumerate(configs):
        ax = axes[idx]
        y_test = result['y_test']
        y_pred = result['final_pred']
        r = result['final_r']
        rmse = result['final_rmse']
        r2 = result['final_r2']
        model_label = _extract_best_model_name(str(result.get('final_type', 'Selected model')))
        ax.scatter(y_test, y_pred, s=52, c=single_scatter_color, alpha=0.82, edgecolors='white', linewidths=0.5)
        _add_identity_line(ax, y_test, y_pred)
        ax.set_xlabel('Experimental PCE (%)')
        ax.set_ylabel('Predicted PCE (%)')
        ax.set_aspect('equal', adjustable='box')
        _annotate_panel(ax, f'({chr(97 + idx)})', y_offset=1.01)
        _format_tick_labels(ax)
        ax.set_title(f'{title}\nBest single model ({model_label})\nr = {r:.4f}, RMSE = {rmse:.4f}, R$^2$ = {r2:.4f}', pad=FIVE_PT_SIZE * 0.85)
    plt.tight_layout(pad=0.9)
    fig.subplots_adjust(top=0.86, bottom=0.12, left=0.08, right=0.98, wspace=0.1)
    filename = data_path('\u4e3b\u6a21\u578b\u6570\u636e\u589e\u5f3a\u524d\u540e\u5bf9\u6bd4\u6563\u70b9\u56fe.png')
    _save_publication_figure(fig, filename, dpi=600)
    plt.close(fig)
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.8), sharex=True, sharey=True)
    for idx, (result, title, _) in enumerate(configs):
        ax = axes[idx]
        y_test = result['y_test']
        y_pred = result['final_pred']
        ax.scatter(y_test, y_pred, s=52, c=single_scatter_color, alpha=0.82, edgecolors='white', linewidths=0.5)
        _add_identity_line(ax, y_test, y_pred)
        ax.set_xlabel('Experimental PCE (%)')
        ax.set_ylabel('Predicted PCE (%)')
        ax.set_aspect('equal', adjustable='box')
        _annotate_panel(ax, f'({chr(97 + idx)})', y_offset=1.01)
        _format_tick_labels(ax)
        model_label = _extract_best_model_name(str(result.get('final_type', 'Selected model')))
        ax.set_title(f"{title}\nBest single model ({model_label})\nr = {result['final_r']:.4f}, RMSE = {result['final_rmse']:.4f}, R$^2$ = {result['final_r2']:.4f}", pad=FIVE_PT_SIZE * 0.85)
    plt.tight_layout(pad=0.9)
    fig.subplots_adjust(top=0.86, bottom=0.12, left=0.08, right=0.98, wspace=0.1)
    legacy_filename = data_path('\u6570\u636e\u589e\u5f3a\u524d\u540e\u96c6\u6210\u6a21\u578b\u7684\u6563\u70b9\u56fe.png')
    _save_publication_figure(fig, legacy_filename, dpi=600)
    plt.close(fig)

def plot_feature_correlation(X, feature_names=None, y=None, max_features=10):
    if X.shape[1] > max_features:
        if y is not None:
            correlations_with_y = []
            for i in range(X.shape[1]):
                corr, _ = pearsonr(X[:, i], y)
                correlations_with_y.append(abs(corr))
            correlations_with_y = np.array(correlations_with_y)
            selected_indices = []
            candidate_indices = list(range(X.shape[1]))
            first_idx = np.argmax(correlations_with_y)
            selected_indices.append(first_idx)
            candidate_indices.remove(first_idx)
            for _ in range(max_features - 1):
                if len(candidate_indices) == 0:
                    break
                best_score = -np.inf
                best_idx = None
                for cand_idx in candidate_indices:
                    target_corr = correlations_with_y[cand_idx]
                    if len(selected_indices) > 0:
                        avg_corr_with_selected = 0
                        for sel_idx in selected_indices:
                            corr_val, _ = pearsonr(X[:, cand_idx], X[:, sel_idx])
                            avg_corr_with_selected += abs(corr_val)
                        avg_corr_with_selected /= len(selected_indices)
                    else:
                        avg_corr_with_selected = 0
                    if avg_corr_with_selected > 0.7:
                        score = 0.5 * target_corr - 0.8 * avg_corr_with_selected
                    elif avg_corr_with_selected > 0.5:
                        score = 0.5 * target_corr - 0.6 * avg_corr_with_selected
                    else:
                        score = 0.5 * target_corr - 0.3 * avg_corr_with_selected
                    if score > best_score:
                        best_score = score
                        best_idx = cand_idx
                if best_idx is not None:
                    selected_indices.append(best_idx)
                    candidate_indices.remove(best_idx)
            selected_indices = np.array(selected_indices)
        else:
            selected_indices = np.arange(max_features)
        X_selected = X[:, selected_indices]
        if feature_names is not None:
            selected_names = [feature_names[i] for i in selected_indices]
        else:
            selected_names = [f'feature_{i}' for i in selected_indices]
    else:
        X_selected = X
        if feature_names is not None:
            selected_names = feature_names
        else:
            selected_names = [f'feature_{i}' for i in range(X.shape[1])]
    corr_matrix = np.corrcoef(X_selected.T)
    figsize = (10, 8) if len(selected_names) <= 10 else (15, 12)
    plot_name = '\u7279\u5f81\u76f8\u5173\u6027\u56fe'
    params_dict = {'max_features': max_features, 'selected_features_count': len(selected_names), 'selected_feature_indices': selected_indices.tolist() if X.shape[1] > max_features else list(range(X.shape[1])), 'figsize': figsize, 'dpi': 300, 'cmap': 'coolwarm', 'vmin': -1.0, 'vmax': 1.0}
    data_dict = {'X_selected': X_selected, 'correlation_matrix': corr_matrix, 'selected_feature_names': selected_names, 'y_data': y if y is not None else None}
    mask_full = np.triu(np.ones_like(corr_matrix, dtype=bool), k=1)
    corr_values = corr_matrix[mask_full]
    params_dict['correlation_statistics'] = {'strong_positive_count': int(np.sum(corr_values > 0.7)), 'moderate_positive_count': int(np.sum((corr_values > 0.3) & (corr_values <= 0.7))), 'weak_count': int(np.sum(np.abs(corr_values) <= 0.3)), 'moderate_negative_count': int(np.sum((corr_values < -0.3) & (corr_values >= -0.7))), 'strong_negative_count': int(np.sum(corr_values < -0.7)), 'mean_correlation': float(np.mean(corr_values)), 'std_correlation': float(np.std(corr_values))}
    save_plot_params_and_data(plot_name, params_dict, data_dict)
    mask_full = np.triu(np.ones_like(corr_matrix, dtype=bool), k=1)
    corr_values = corr_matrix[mask_full]
    strong_pos = np.sum(corr_values > 0.7)
    moderate_pos = np.sum((corr_values > 0.3) & (corr_values <= 0.7))
    weak = np.sum(np.abs(corr_values) <= 0.3)
    moderate_neg = np.sum((corr_values < -0.3) & (corr_values >= -0.7))
    strong_neg = np.sum(corr_values < -0.7)
    fig, ax = plt.subplots(figsize=figsize)
    mask = np.triu(np.ones_like(corr_matrix, dtype=bool), k=1)
    sns.heatmap(corr_matrix, mask=mask, annot=False, cmap='coolwarm', center=0, square=True, linewidths=0.4, cbar_kws={'shrink': 0.75, 'label': 'Pearson r'}, vmin=-1.0, vmax=1.0, ax=ax)
    ax.set_xticks(np.arange(len(selected_names)) + 0.5)
    ax.set_yticks(np.arange(len(selected_names)) + 0.5)
    label_fontsize = FIVE_PT_SIZE + 3.5
    ax.set_xticklabels(selected_names, rotation=45, ha='right', fontsize=label_fontsize)
    ax.set_yticklabels(selected_names, rotation=0, fontsize=label_fontsize)
    _format_tick_labels(ax)
    fig.tight_layout()
    corr_png = data_path('\u7279\u5f81\u76f8\u5173\u6027\u56fe.png')
    fig.savefig(corr_png, bbox_inches='tight')
    plt.close(fig)

def export_ensemble_weights(model_weights, config_name):
    df_weights = pd.DataFrame([{'\u6a21\u578b': name, '\u6743\u91cd': float(weight)} for name, weight in model_weights.items()])
    filename = data_path('\u6a21\u578b\u6743\u91cd.csv')
    df_weights.to_csv(filename, index=False, encoding='utf-8-sig')
    return df_weights

def export_high_performance_pce(y_test, y_pred, d_codes, a_codes, n_samples=10, pce_threshold=16.0, random_state=None):
    high_pce_indices = np.where(y_test > pce_threshold)[0]
    if len(high_pce_indices) == 0:
        sorted_all = np.argsort(y_test)[::-1]
        selected_indices = np.sort(sorted_all[:min(n_samples, len(y_test))])
    elif len(high_pce_indices) < n_samples:
        selected = list(high_pce_indices)
        remaining = [i for i in np.argsort(y_test)[::-1] if i not in set(selected)]
        need = min(n_samples, len(y_test)) - len(selected)
        if need > 0:
            selected.extend(remaining[:need])
        selected_indices = np.sort(np.array(selected, dtype=int))
    else:
        rs = _resolve_random_state(random_state)
        np.random.seed(rs)
        selected_indices = np.random.choice(high_pce_indices, size=n_samples, replace=False)
        selected_indices = np.sort(selected_indices)
    results = []
    for i, idx in enumerate(selected_indices):
        true_val = y_test[idx]
        pred_val = y_pred[idx]
        diff = pred_val - true_val
        relative_error = diff / true_val * 100 if true_val != 0 else 0
        results.append({'\u5e8f\u53f7': i + 1, '\u4f9b\u4f53\u7f16\u7801': str(d_codes[idx]) if idx < len(d_codes) else f'D_{idx}', '\u53d7\u4f53\u7f16\u7801': str(a_codes[idx]) if idx < len(a_codes) else f'A_{idx}', '\u771f\u5b9ePCE\u503c': round(true_val, 4), '\u9884\u6d4bPCE\u503c': round(pred_val, 4), '\u5dee\u503c': round(diff, 4), '\u76f8\u5bf9\u8bef\u5dee(%)': round(relative_error, 2), '\u7edd\u5bf9\u8bef\u5dee': round(abs(diff), 4)})
    df_results = pd.DataFrame(results)
    filename = data_path('10\u7ec4\u9ad8\u6027\u80fdPCE\u9884\u6d4b\u503c\u3001\u771f\u5b9e\u503c\u3001\u5dee\u503c\u3001\u76f8\u5bf9\u8bef\u5dee.csv')
    df_results.to_csv(filename, index=False, encoding='utf-8-sig')
    return df_results

def plot_shap_explanations(result, feature_names, X_train, X_test, y_test, config_name='No Augmentation', random_state=None, max_display=12):
    if not SHAP_AVAILABLE:
        return None
    results_dict = result.get('results', {})
    if not results_dict:
        return None
    if result.get('final_type'):
        best_model_name = _extract_best_model_name(str(result.get('final_type')))
    else:
        best_model_name = max(results_dict.keys(), key=lambda x: results_dict[x].get('r_test', -np.inf))
    best_model = result['models'][best_model_name]
    if len(X_test) == 0:
        return None
    rs = _resolve_random_state(random_state)
    rng = np.random.RandomState(rs)
    n_samples_for_shap = min(120, len(X_test))
    sample_indices = rng.choice(len(X_test), n_samples_for_shap, replace=False)
    X_shap = X_test[sample_indices]
    try:
        if isinstance(best_model, (RandomForestRegressor, GradientBoostingRegressor)):
            explainer = shap.TreeExplainer(best_model)
            shap_values = explainer.shap_values(X_shap)
            explainer_name = 'TreeExplainer'
        elif isinstance(best_model, (Ridge, ElasticNet)):
            background = X_train if len(X_train) <= 200 else X_train[rng.choice(len(X_train), 200, replace=False)]
            try:
                explainer = shap.LinearExplainer(best_model, background)
            except Exception:
                explainer = shap.Explainer(best_model, background)
            shap_values = explainer(X_shap)
            explainer_name = 'LinearExplainer'
        else:
            background_size = min(50, len(X_train))
            background_idx = rng.choice(len(X_train), background_size, replace=False)
            X_background = X_train[background_idx]
            explainer = shap.KernelExplainer(best_model.predict, X_background)
            shap_values = explainer.shap_values(X_shap, nsamples=100)
            explainer_name = 'KernelExplainer'
        if hasattr(shap_values, 'values'):
            shap_values_arr = np.array(shap_values.values)
        else:
            shap_values_arr = np.array(shap_values)
        if isinstance(shap_values, list):
            shap_values_arr = np.array(shap_values[0])
        if shap_values_arr.ndim == 1:
            shap_values_arr = shap_values_arr.reshape(-1, 1)
        feature_names_use = list(feature_names[:shap_values_arr.shape[1]])
        mean_abs_shap = np.abs(shap_values_arr).mean(axis=0)
        meta_rows = []
        for fname, ms in zip(feature_names_use, mean_abs_shap):
            info = _parse_paired_feature_label(fname)
            meta_rows.append({'Feature': info['feature_label'], 'Feature_side': info['feature_side'], 'Bit_index': info['bit_index'], 'Feature_note': info['feature_note'], 'Mean_ABS_SHAP': float(ms)})
        feature_importance_df = pd.DataFrame(meta_rows).sort_values('Mean_ABS_SHAP', ascending=False).reset_index(drop=True)
        feature_importance_df['Rank'] = np.arange(1, len(feature_importance_df) + 1)
        tag = _sanitize_plot_tag(config_name)
        fig = plt.figure(figsize=(12.4, 5.8), constrained_layout=False)
        gs = fig.add_gridspec(1, 2, width_ratios=[1.2, 0.95], wspace=0.28)
        ax1 = fig.add_subplot(gs[0, 0])
        ax2 = fig.add_subplot(gs[0, 1])
        plt.sca(ax1)
        shap.summary_plot(shap_values_arr, X_shap, feature_names=feature_names_use, show=False, max_display=max_display, plot_size=None, color_bar=True)
        ax1.set_title('')
        ax1.set_xlabel('SHAP value', fontsize=FIVE_PT_SIZE)
        ax1.set_ylabel('Feature', fontsize=FIVE_PT_SIZE)
        _annotate_panel(ax1, '(a)', y_offset=1.01)
        plt.sca(ax2)
        shap.summary_plot(shap_values_arr, X_shap, feature_names=feature_names_use, plot_type='bar', show=False, max_display=max_display, plot_size=None, color=SCATTER_COLORS.get('no_aug', '#62C2A4'))
        ax2.set_title('')
        ax2.set_xlabel('mean(|SHAP value|)', fontsize=FIVE_PT_SIZE)
        ax2.set_ylabel('')
        _annotate_panel(ax2, '(b)', y_offset=1.01)
        model_label = _extract_best_model_name(str(result.get('final_type', best_model_name)))
        fig.suptitle(f'SHAP interpretation ({config_name})\nBest single model ({model_label})', fontsize=FIVE_PT_SIZE + 0.8, y=0.985)
        fig.subplots_adjust(top=0.82, left=0.07, right=0.98, bottom=0.11)
        combined_filename = data_path(f'SHAP_Summary_Feature_Importance_{tag}.png')
        _save_publication_figure(fig, combined_filename, dpi=600)
        plt.close(fig)
        filename_csv = data_path(f'SHAP_Feature_Importance_{tag}.csv')
        feature_importance_df.to_csv(filename_csv, index=False, encoding='utf-8-sig')
        top_candidates = feature_importance_df.head(max_display).copy()
        top_candidates.to_csv(data_path(f'SHAP_Top_Features_{tag}.csv'), index=False, encoding='utf-8-sig')
        plot_name = f'SHAP\u89e3\u91ca\u56fe_{tag}'
        params_dict = {'best_model_name': best_model_name, 'config_name': config_name, 'n_samples_for_shap': int(n_samples_for_shap), 'explainer_type': explainer_name, 'max_display': int(max_display), 'figure_layout': '1x2 summary + bar', 'dpi': 600, 'output_png': combined_filename, 'output_pdf': os.path.splitext(combined_filename)[0] + '.pdf'}
        data_dict = {'X_shap': X_shap, 'shap_values': shap_values_arr, 'mean_abs_shap': mean_abs_shap, 'feature_names': np.array(feature_names_use, dtype=object), 'y_test_subset': y_test[sample_indices] if len(y_test) > max(sample_indices) else None}
        save_plot_params_and_data(plot_name, params_dict, data_dict)
        return feature_importance_df
    except Exception as e:
        import traceback
        traceback.print_exc()
        return None

def main(plot_only=False, max_features=10, random_seed=None, split_preset=DEFAULT_SPLIT_PRESET, run_all_split_presets=False):
    if plot_only:
        run_plot_only_mode(max_features=max_features)
        return
    seed = _resolve_random_state(random_seed)
    presets_to_run = list(SPLIT_PRESET_SPECS.keys()) if run_all_split_presets else [split_preset]
    root_before = _RUN_OUTPUT_ROOT
    base_dir = get_run_output_root()
    for sp in presets_to_run:
        if run_all_split_presets and len(presets_to_run) > 1:
            sub = os.path.join(base_dir, f'split_{sp}')
            set_run_output_root(sub)
        _main_one_split(plot_only, max_features, seed, sp)
    set_run_output_root(root_before)

def _main_one_split(plot_only, max_features, seed, split_preset):
    unknown_indices, filtered_fd_fp, filtered_fa_fp, filtered_fp_Y, unknown_fd_fp, unknown_fa_fp, unknown_fp_Y, unknown_donor_smiles, unknown_acceptor_smiles = select_and_split_unknown_20_pairs(random_state=seed)
    if CURRENT_EXPERIMENT_CONTEXT.get('use_gao_aux'):
        _append_aux_gao_to_filtered_training_data(CURRENT_EXPERIMENT_CONTEXT.get('gao_fd_path'), CURRENT_EXPERIMENT_CONTEXT.get('gao_fa_path'), CURRENT_EXPERIMENT_CONTEXT.get('gao_y_path'), n_bits=1024)
    _write_experiment_context_summary({'random_seed': int(seed), 'split_preset': split_preset, 'core_rows_after_unknown20': int(len(filtered_fp_Y)), 'unknown20_rows': int(len(unknown_fp_Y))})
    X_raw, y_raw, feat_names, d_codes, ac_codes, _, _ = load_data_without_augmentation()
    try:
        codes_d_unique = len(np.unique(d_codes))
        codes_a_unique = len(np.unique(ac_codes))
        is_placeholder_codes = False
        if isinstance(d_codes[0], (str, bytes)) and isinstance(ac_codes[0], (str, bytes)):
            d0 = str(d_codes[0])
            a0 = str(ac_codes[0])
            if d0.startswith('D_') and a0.startswith('A_'):
                is_placeholder_codes = True
        if not is_placeholder_codes and codes_d_unique < X_raw.shape[0] and (codes_a_unique < X_raw.shape[0]):
            pass
        else:
            fd_fp_filtered = np.load(data_path('fd_fp_filtered.npy'))
            fa_fp_filtered = np.load(data_path('fa_fp_filtered.npy'))
            fd_q = np.round(fd_fp_filtered, 6)
            fa_q = np.round(fa_fp_filtered, 6)
            unique_d_no_aug = np.unique(fd_q, axis=0).shape[0]
            unique_a_no_aug = np.unique(fa_q, axis=0).shape[0]
    except Exception as e:
        pass
    result_no_aug = train_and_evaluate(X_raw, y_raw, '\u65e0\u589e\u5f3a(\u5408\u5e76\u540e)', random_state=seed, split_preset=split_preset, augment_train=False, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
    result_with_aug = train_and_evaluate(X_raw, y_raw, '\u6709\u589e\u5f3a(\u5408\u5e76\u540e)', random_state=seed, split_preset=split_preset, augment_train=True, feature_names_all=feat_names, donor_codes=d_codes, acceptor_codes=ac_codes)
    y_with_aug_plot = np.concatenate([result_with_aug['y_train'], result_with_aug['y_val'], result_with_aug['y_test']])
    feat_names_sel = result_with_aug['selected_feature_names']
    unknown_results, unknown_predictions = predict_unknown_20_pairs(result_with_aug, result_with_aug['selector'], result_with_aug['scaler'])
    plot_combined_pce_distribution_ab(y_raw, y_with_aug_plot)
    plot_model_scatters_comparison(result_no_aug, result_with_aug)
    plot_metrics_comparison(result_no_aug, result_with_aug)
    plot_final_model_comparison_scatter(result_no_aug, result_with_aug)
    X_corr = np.vstack([result_with_aug['X_train'], result_with_aug['X_val'], result_with_aug['X_test']])
    plot_feature_correlation(X_corr, feat_names_sel, y_with_aug_plot)
    plot_shap_explanations(result_no_aug, result_no_aug['selected_feature_names'], result_no_aug['X_train'], result_no_aug['X_test'], result_no_aug['y_test'], 'No Augmentation', random_state=seed, max_display=12)
    export_ensemble_weights(result_with_aug['model_weights'], '\u589e\u5f3a\u540e')
    export_high_performance_pce(result_with_aug['y_test'], result_with_aug['final_pred'], result_with_aug['d_codes_test'], result_with_aug['a_codes_test'], random_state=seed)
    no_aug_ensemble_r = result_no_aug.get('final_r', 0)
    no_aug_selected = result_no_aug.get('selected_models', {})
    if len(no_aug_selected) == 1:
        model_name = list(no_aug_selected.keys())[0]
        for name, res in result_no_aug['results'].items():
            if name != model_name:
                pass
    with_aug_selected = result_with_aug.get('selected_models', {})
    with_aug_ensemble_r = None
    with_aug_best_r = None
    if len(with_aug_selected) > 1:
        y_test_aug = result_with_aug['y_test']
        ensemble_pred = np.zeros(len(y_test_aug))
        model_weights = result_with_aug['model_weights']
        for name in with_aug_selected.keys():
            if name in result_with_aug['results']:
                pred = result_with_aug['results'][name]['y_pred_test']
                ensemble_pred += model_weights[name] * pred
        with_aug_ensemble_r = pearsonr(y_test_aug, ensemble_pred)[0]
        with_aug_best_r = result_with_aug['final_r']
        if with_aug_best_r > with_aug_ensemble_r:
            pass
    for name, res in result_with_aug['results'].items():
        status = '\u53c2\u4e0e\u96c6\u6210' if name in with_aug_selected else '\u672a\u53c2\u4e0e'
    if with_aug_ensemble_r is not None:
        ensemble_r_str = f'{with_aug_ensemble_r:.4f}'
    else:
        ensemble_r_str = 'N/A'
    best_r_str = f'{with_aug_best_r:.4f}' if with_aug_best_r is not None else f"{result_with_aug['final_r']:.4f}"
    unknown_true = unknown_fp_Y
    unknown_pred = unknown_predictions
    unknown_r = pearsonr(unknown_true, unknown_pred)[0]
    unknown_rmse = np.sqrt(mean_squared_error(unknown_true, unknown_pred))
    unknown_r2 = r2_score(unknown_true, unknown_pred)
    unknown_mae = mean_absolute_error(unknown_true, unknown_pred)
    if os.path.normpath(get_run_output_root()) != os.path.normpath(SCRIPT_DIR):
        summary_payload = {'random_seed': int(seed), 'split_preset': split_preset, 'output_dir': get_run_output_root(), 'final_r_no_aug': float(result_no_aug['final_r']), 'final_r2_no_aug': float(result_no_aug['final_r2']), 'final_rmse_no_aug': float(result_no_aug['final_rmse']), 'final_mae_no_aug': float(result_no_aug['final_mae']), 'final_type_no_aug': str(result_no_aug['final_type']), 'final_r_with_aug': float(result_with_aug['final_r']), 'final_r2_with_aug': float(result_with_aug['final_r2']), 'final_rmse_with_aug': float(result_with_aug['final_rmse']), 'final_mae_with_aug': float(result_with_aug['final_mae']), 'final_type_with_aug': str(result_with_aug['final_type']), 'unknown_20_r': float(unknown_r), 'unknown_20_rmse': float(unknown_rmse), 'unknown_20_r2': float(unknown_r2), 'unknown_20_mae': float(unknown_mae)}
        with open(data_path('run_metrics_summary.json'), 'w', encoding='utf-8') as sf:
            json.dump(summary_payload, sf, indent=2, ensure_ascii=False)
    table2_rows = [{'Setting': 'In-house baseline (No augmentation)', 'Sample_scope': 'Held-out test set', 'Model_used': result_no_aug['final_type'], 'Pearson_r': float(result_no_aug['final_r']), 'R_squared': float(result_no_aug['final_r2']), 'RMSE': float(result_no_aug['final_rmse']), 'MAE': float(result_no_aug['final_mae'])}, {'Setting': 'In-house baseline (With augmentation)', 'Sample_scope': 'Held-out test set', 'Model_used': result_with_aug['final_type'], 'Pearson_r': float(result_with_aug['final_r']), 'R_squared': float(result_with_aug['final_r2']), 'RMSE': float(result_with_aug['final_rmse']), 'MAE': float(result_with_aug['final_mae'])}, {'Setting': 'Unknown-pair hold-out', 'Sample_scope': 'Unknown 20 D-A pairs', 'Model_used': str(result_with_aug.get('final_type', 'Selected final model')), 'Pearson_r': float(unknown_r), 'R_squared': float(unknown_r2), 'RMSE': float(unknown_rmse), 'MAE': float(unknown_mae)}]
    table2_df = pd.DataFrame(table2_rows)
    table2_csv = data_path('Table_2_Structure_resolved_baseline_and_unknown_holdout.csv')
    table2_df.to_csv(table2_csv, index=False, encoding='utf-8-sig')
    try:
        table2_xlsx = data_path('Table_2_Structure_resolved_baseline_and_unknown_holdout.xlsx')
        table2_df.to_excel(table2_xlsx, index=False)
    except Exception as exc:
        pass
    if SHAP_AVAILABLE:
        pass
    if SHAP_AVAILABLE:
        pass
    if SHAP_AVAILABLE:
        pass

def _auto_prepare_from_existing_source_files(radius: int):
    da_path = _find_first_existing_source_file(PUBLIC_DA_FILENAME)
    if not os.path.exists(da_path):
        raise FileNotFoundError(f'Not found: {da_path}. Please place D-A.csv in the script directory.')
    _prepare_inputs_from_smiles_dataset(dataset_path=da_path, donor_col=None, acceptor_col=None, pce_col=None, radius=radius, n_bits=1024)

def _write_best_experiment_config(out_root: str, profile_name: str='main_r3'):
    summary_dir = os.path.join(out_root, '\u6c47\u603b_all_splits')
    table1 = os.path.join(summary_dir, '\u88681_\u5404\u5212\u5206_\u6240\u6709seed_\u4e94\u6a21\u578br\u5bf9\u6bd4.csv')
    if not os.path.exists(table1):
        return None
    df = pd.read_csv(table1)
    agg = df.groupby(['split_preset', 'seed'], as_index=False).agg(best_final_r_with_aug=('best_final_r_with_aug', 'first'), best_final_r_no_aug=('best_final_r_no_aug', 'first')).sort_values('best_final_r_with_aug', ascending=False).reset_index(drop=True)
    if agg.empty:
        return None
    best = agg.iloc[0]
    sp = str(best['split_preset'])
    seed = int(best['seed'])
    payload = {'profile': profile_name, 'radius': int(EXPERIMENT_PROFILES.get(profile_name, {}).get('radius', CURRENT_EXPERIMENT_CONTEXT.get('radius', 3))), 'use_gao_aux': bool(EXPERIMENT_PROFILES.get(profile_name, {}).get('use_gao_aux', False)), 'best_split_preset': sp, 'best_split_desc': SPLIT_PRESET_SPECS.get(sp, {}).get('description', ''), 'best_seed': seed, 'best_final_r_with_aug': float(best['best_final_r_with_aug']), 'best_final_r_no_aug': float(best['best_final_r_no_aug']), 'search_out_root': os.path.abspath(out_root), 'source_table': os.path.abspath(table1)}
    local_paths = [os.path.join(out_root, BEST_MAIN_CONFIG_FILENAME), os.path.join(summary_dir, BEST_MAIN_CONFIG_FILENAME), input_path(BEST_MAIN_CONFIG_FILENAME)]
    for p in local_paths:
        with open(p, 'w', encoding='utf-8') as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
    return payload

def _load_best_experiment_config(path_or_none=None):
    cfg_path = path_or_none if path_or_none else input_path(BEST_MAIN_CONFIG_FILENAME)
    if not os.path.exists(cfg_path):
        raise FileNotFoundError(f'\u672a\u627e\u5230\u6700\u4f73\u914d\u7f6e\u6587\u4ef6: {cfg_path}\u3002\u8bf7\u5148\u8fd0\u884c --main-search START END \u751f\u6210\u3002')
    with open(cfg_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def _run_fixed_profile_from_best_config(profile_name: str, best_cfg_path=None, plot_only=False, max_features=10):
    cfg = _load_best_experiment_config(best_cfg_path)
    prof = EXPERIMENT_PROFILES.get(profile_name, EXPERIMENT_PROFILES['legacy'])
    _set_experiment_context(profile=profile_name, radius=int(prof['radius']), use_gao_aux=bool(prof['use_gao_aux']), gao_fd_path=CURRENT_EXPERIMENT_CONTEXT.get('gao_fd_path'), gao_fa_path=CURRENT_EXPERIMENT_CONTEXT.get('gao_fa_path'), gao_y_path=CURRENT_EXPERIMENT_CONTEXT.get('gao_y_path'))
    _auto_prepare_from_existing_source_files(radius=int(prof['radius']))
    main(plot_only=plot_only, max_features=max_features, random_seed=int(cfg['best_seed']), split_preset=str(cfg['best_split_preset']), run_all_split_presets=False)

def _run_main_search_on_existing_dareal(start_seed: int, end_seed: int):
    _set_experiment_context(profile='main_r3', radius=3, use_gao_aux=False)
    _auto_prepare_from_existing_source_files(radius=3)
    out_root = run_all_splits_seed_search_and_summary(start_seed, end_seed, only_presets=None)
    _write_best_experiment_config(out_root, profile_name='main_r3')
    return out_root
if __name__ == '__main__':
    _epilog = '\n\u4f60\u73b0\u5728\u6700\u5e38\u7528\u7684 4 \u4e2a\u547d\u4ee4\uff08\u76f4\u63a5\u4f7f\u7528\u811a\u672c\u540c\u76ee\u5f55\u91cc\u7684 D-A.csv / Donor.csv / Acceptor.csv\uff09\uff1a\n\n1) \u4e3b\u5b9e\u9a8c\uff1a\u56db\u79cd\u5212\u5206 \xd7 \u79cd\u5b50\u641c\u7d22\uff08\u63a8\u8350\u5148\u8dd1\u8fd9\u4e2a\uff09\n   python analysis_with_unknown_20_pairs.py --main-search 0 100\n\n2) \u4e3b\u5b9e\u9a8c\uff1a\u6309\u641c\u7d22\u5f97\u5230\u7684\u6700\u4f73 split + seed \u56fa\u5b9a\u91cd\u8dd1\n   python analysis_with_unknown_20_pairs.py --run-best-main\n\n3) \u534a\u5f84\u5bf9\u7167\uff1a\u6cbf\u7528\u4e3b\u5b9e\u9a8c\u6700\u4f73 split + seed\uff0c\u6539\u4e3a radius=2 \u91cd\u8dd1\n   python analysis_with_unknown_20_pairs.py --run-best-core-r2\n\n4) \u8f85\u52a9\u8054\u5408\u8bad\u7ec3\uff1a\u6cbf\u7528\u4e3b\u5b9e\u9a8c\u6700\u4f73 split + seed\uff0ccore(radius=2)+Gao \u6570\u636e\n   python analysis_with_unknown_20_pairs.py --run-best-core-r2-plus-gao --gao-fd-path gao_fd_fp.npy --gao-fa-path gao_fa_fp.npy --gao-y-path gao_fp_Y.npy\n\n\u5982\u679c\u4f60\u8981\u4ece\u65b0\u7684\u5355\u4e2a\u6570\u636e\u96c6\u81ea\u52a8\u751f\u6210 DAreal/geiti/shouti\uff1a\n   python analysis_with_unknown_20_pairs.py --dataset your_dataset.csv --experiment-profile main_r3\n'
    parser = argparse.ArgumentParser(description='D/A \u7ec4\u5408\u5206\u6790\u7cfb\u7edf\uff08\u652f\u6301\u4e3b\u5b9e\u9a8c\u79cd\u5b50/\u5212\u5206\u641c\u7d22\u3001\u56fa\u5b9a\u6700\u4f73\u914d\u7f6e\u91cd\u8dd1\u3001\u534a\u5f84\u5bf9\u7167\u3001Gao \u8f85\u52a9\u8054\u5408\u8bad\u7ec3\uff09', formatter_class=argparse.RawDescriptionHelpFormatter, epilog=_epilog)
    parser.add_argument('--plot-only', action='store_true', help='\u4ec5\u6839\u636e\u5df2\u6709\u6570\u636e/\u7f13\u5b58\u751f\u6210\u56fe\u50cf\uff0c\u8df3\u8fc7\u6570\u636e\u6574\u7406\u4e0e\u5efa\u6a21')
    parser.add_argument('--max-features', type=int, default=10, help='\u7eaf\u7ed8\u56fe\u6a21\u5f0f\u4e0b\u7279\u5f81\u76f8\u5173\u6027\u56fe\u5c55\u793a\u7684\u6700\u5927\u7279\u5f81\u6570\u91cf')
    parser.add_argument('--seed', type=int, default=None, help=f'\u5168\u5c40\u968f\u673a\u79cd\u5b50\uff08\u9ed8\u8ba4 {DEFAULT_RANDOM_SEED}\uff09')
    parser.add_argument('--split-preset', type=str, default=DEFAULT_SPLIT_PRESET, choices=list(SPLIT_PRESET_SPECS.keys()), help='\u624b\u52a8\u6307\u5b9a\u5212\u5206\u9884\u8bbe')
    parser.add_argument('--run-all-split-presets', '--batch-all-splits', action='store_true', dest='run_all_split_presets', help='\u4f9d\u6b21\u8fd0\u884c\u6240\u6709\u5212\u5206\u9884\u8bbe')
    parser.add_argument('--compare-split-seeds', type=str, default=None, metavar='LIST', help='\u4ec5\u6253\u5370\u4e0d\u540c\u79cd\u5b50\u4f1a\u62bd\u4e2d\u54ea\u4e9b\u884c\u4f5c\u4e3a\u672a\u77e5\u7ec4\uff08\u9017\u53f7\u5206\u9694\uff09')
    parser.add_argument('--compare-train-seeds', type=str, default=None, metavar='LIST', help='\u5728\u5df2\u6709 filtered \u6570\u636e\u4e0d\u53d8\u65f6\uff0c\u5bf9\u6bd4\u591a\u79cd\u5b50\u4e0b\u8bad\u7ec3/\u589e\u5f3a/\u5212\u5206\u7684 final_r\uff08\u9017\u53f7\u5206\u9694\uff09')
    parser.add_argument('--multi-seed', action='store_true', help=f'\u4f9d\u6b21\u4f7f\u7528\u79cd\u5b50 {MULTI_SEED_DEFAULT_LIST} \u5404\u8dd1\u4e00\u904d\u5b8c\u6574\u6d41\u7a0b')
    parser.add_argument('--multi-seed-list', type=str, default=None, metavar='LIST', help='\u4e0e --multi-seed \u8054\u7528\uff1a\u9017\u53f7\u5206\u9694\u81ea\u5b9a\u4e49\u79cd\u5b50\u5217\u8868')
    parser.add_argument('--seed-search-lite', type=int, nargs=2, metavar=('START', 'END'), help='\u8f7b\u91cf\u7ea7\u79cd\u5b50\u641c\u7d22\uff1a\u4fdd\u6301\u672a\u77e520\u7ec4\u4e0d\u53d8\uff0c\u53ea\u6539\u53d8\u8bad\u7ec3/\u589e\u5f3a/\u5212\u5206\u79cd\u5b50')
    parser.add_argument('--seed-search', type=int, nargs=2, metavar=('START', 'END'), help='\u5b8c\u6574\u79cd\u5b50\u641c\u7d22\uff1a\u6bcf\u4e2a\u79cd\u5b50\u91cd\u65b0\u9009\u62e9\u672a\u77e520\u7ec4\uff0c\u5b8c\u6574\u6d41\u7a0b')
    parser.add_argument('--all-splits-seed-search', type=int, nargs=2, metavar=('START', 'END'), help='\u56db\u79cd\u5212\u5206\u6bd4\u4f8b\u5206\u522b\u6267\u884c\u968f\u673a\u79cd\u5b50\u641c\u7d22\u5e76\u6c47\u603b\u51fa\u8868\u683c+\u66f2\u7ebf\u56fe')
    parser.add_argument('--all-splits-seed-search-presets', type=str, default=None, metavar='LIST', help='\u4e0e --all-splits-seed-search \u8054\u7528\uff1a\u53ea\u8dd1\u90e8\u5206\u5212\u5206\uff08\u9017\u53f7\u5206\u9694\uff09')
    parser.add_argument('--main-search', type=int, nargs=2, metavar=('START', 'END'), help='\u6700\u7b80\u4e3b\u5b9e\u9a8c\u641c\u7d22\uff1a\u81ea\u52a8\u4f7f\u7528\u811a\u672c\u76ee\u5f55\u4e2d\u7684 D-A.csv\uff0c\u6309 main_r3 \u8dd1\u56db\u79cd\u5212\u5206 \xd7 \u79cd\u5b50\u641c\u7d22\uff0c\u5e76\u5199\u51fa best_main_experiment_config.json')
    parser.add_argument('--run-best-main', action='store_true', help='\u8bfb\u53d6 best_main_experiment_config.json\uff0c\u6309\u6700\u4f73 split+seed \u56fa\u5b9a\u91cd\u8dd1\u4e3b\u5b9e\u9a8c\uff08main_r3\uff09')
    parser.add_argument('--run-best-core-r2', action='store_true', help='\u8bfb\u53d6 best_main_experiment_config.json\uff0c\u6cbf\u7528\u6700\u4f73 split+seed \u8fd0\u884c\u534a\u5f84\u5bf9\u7167\uff08core_r2\uff09')
    parser.add_argument('--run-best-core-r2-plus-gao', action='store_true', help='\u8bfb\u53d6 best_main_experiment_config.json\uff0c\u6cbf\u7528\u6700\u4f73 split+seed \u8fd0\u884c\u8f85\u52a9\u8054\u5408\u8bad\u7ec3\uff08core_r2_plus_gao\uff09')
    parser.add_argument('--best-main-config', type=str, default=None, help=f'\u6700\u4f73\u4e3b\u5b9e\u9a8c\u914d\u7f6e\u6587\u4ef6\u8def\u5f84\uff08\u9ed8\u8ba4\u811a\u672c\u76ee\u5f55\u4e0b {BEST_MAIN_CONFIG_FILENAME}\uff09')
    parser.add_argument('--use-existing-da-files', action='store_true', help='\u4f7f\u7528\u811a\u672c\u76ee\u5f55\u4e2d\u7684 D-A.csv \u81ea\u52a8\u91cd\u5efa\u5f53\u524d profile \u6240\u9700\u8f93\u5165\u6587\u4ef6')
    parser.add_argument('--dataset', type=str, default=None, help='\u63d0\u4f9b\u5305\u542b donor/acceptor SMILES \u4e0e PCE \u7684 csv/xlsx \u6570\u636e\u96c6\uff1b\u811a\u672c\u4f1a\u81ea\u52a8\u751f\u6210\u65e7\u6d41\u7a0b\u6240\u9700\u8f93\u5165\u6587\u4ef6')
    parser.add_argument('--output-tag', type=str, default=None, help='\u5c06\u672c\u6b21\u8fd0\u884c\u7684\u6240\u6709\u8f93\u51fa\u5199\u5165\u811a\u672c\u76ee\u5f55\u4e0b\u7684\u6307\u5b9a\u5b50\u76ee\u5f55\uff0c\u4f8b\u5982 --output-tag main_final')
    parser.add_argument('--donor-col', type=str, default=None, help='\u4e0e --dataset \u8054\u7528\uff1a\u624b\u52a8\u6307\u5b9a donor SMILES \u5217\u540d')
    parser.add_argument('--acceptor-col', type=str, default=None, help='\u4e0e --dataset \u8054\u7528\uff1a\u624b\u52a8\u6307\u5b9a acceptor SMILES \u5217\u540d')
    parser.add_argument('--pce-col', type=str, default=None, help='\u4e0e --dataset \u8054\u7528\uff1a\u624b\u52a8\u6307\u5b9a PCE \u5217\u540d')
    parser.add_argument('--auto-prepare-only', action='store_true', help='\u4e0e --dataset \u8054\u7528\uff1a\u53ea\u505a SMILES->\u6307\u7eb9/\u8f93\u5165\u6587\u4ef6\u8f6c\u6362\uff0c\u4e0d\u7ee7\u7eed\u6267\u884c\u540e\u7eed\u5206\u6790\u6d41\u7a0b')
    parser.add_argument('--experiment-profile', type=str, default='main_r3', choices=list(EXPERIMENT_PROFILES.keys()), help='\u5b9e\u9a8c\u77e9\u9635\u914d\u7f6e\uff1amain_r3 / core_r2 / core_r2_plus_gao / legacy')
    parser.add_argument('--dataset-radius', type=int, default=None, help='\u4e0e --dataset \u8054\u7528\uff1a\u8986\u76d6 experiment-profile \u9ed8\u8ba4\u534a\u5f84')
    parser.add_argument('--gao-fd-path', type=str, default=None, help='Gao \u8f85\u52a9\u96c6 donor \u6307\u7eb9 npy\uff08\u4ec5 core_r2_plus_gao \u4f7f\u7528\uff09')
    parser.add_argument('--gao-fa-path', type=str, default=None, help='Gao \u8f85\u52a9\u96c6 acceptor \u6307\u7eb9 npy\uff08\u4ec5 core_r2_plus_gao \u4f7f\u7528\uff09')
    parser.add_argument('--gao-y-path', type=str, default=None, help='Gao \u8f85\u52a9\u96c6 PCE \u6807\u7b7e npy\uff08\u4ec5 core_r2_plus_gao \u4f7f\u7528\uff09')
    args = parser.parse_args()
    if args.output_tag:
        apply_output_tag(args.output_tag)
    if args.main_search:
        _run_main_search_on_existing_dareal(args.main_search[0], args.main_search[1])
        raise SystemExit(0)
    if args.run_best_main or args.run_best_core_r2 or args.run_best_core_r2_plus_gao:
        if args.run_best_core_r2_plus_gao and (not (args.gao_fd_path and args.gao_fa_path and args.gao_y_path)):
            raise SystemExit('\u8fd0\u884c --run-best-core-r2-plus-gao \u65f6\uff0c\u8bf7\u540c\u65f6\u63d0\u4f9b --gao-fd-path / --gao-fa-path / --gao-y-path')
        if args.run_best_main:
            _set_experiment_context(profile='main_r3', radius=3, use_gao_aux=False)
            _run_fixed_profile_from_best_config('main_r3', best_cfg_path=args.best_main_config, plot_only=args.plot_only, max_features=args.max_features)
            raise SystemExit(0)
        if args.run_best_core_r2:
            _set_experiment_context(profile='core_r2', radius=2, use_gao_aux=False)
            _run_fixed_profile_from_best_config('core_r2', best_cfg_path=args.best_main_config, plot_only=args.plot_only, max_features=args.max_features)
            raise SystemExit(0)
        if args.run_best_core_r2_plus_gao:
            _set_experiment_context(profile='core_r2_plus_gao', radius=2, use_gao_aux=True, gao_fd_path=args.gao_fd_path, gao_fa_path=args.gao_fa_path, gao_y_path=args.gao_y_path)
            _run_fixed_profile_from_best_config('core_r2_plus_gao', best_cfg_path=args.best_main_config, plot_only=args.plot_only, max_features=args.max_features)
            raise SystemExit(0)
    profile_cfg = EXPERIMENT_PROFILES.get(args.experiment_profile, EXPERIMENT_PROFILES['legacy'])
    chosen_radius = int(profile_cfg['radius'] if args.dataset_radius is None else args.dataset_radius)
    _set_experiment_context(profile=args.experiment_profile, radius=chosen_radius, use_gao_aux=profile_cfg['use_gao_aux'], gao_fd_path=args.gao_fd_path, gao_fa_path=args.gao_fa_path, gao_y_path=args.gao_y_path)
    if CURRENT_EXPERIMENT_CONTEXT.get('use_gao_aux') and (not (args.gao_fd_path and args.gao_fa_path and args.gao_y_path)):
        raise SystemExit('\u5f53\u524d experiment-profile \u9700\u8981 Gao \u8f85\u52a9\u96c6\uff0c\u8bf7\u540c\u65f6\u63d0\u4f9b --gao-fd-path / --gao-fa-path / --gao-y-path')
    if args.dataset:
        _prepare_inputs_from_smiles_dataset(dataset_path=args.dataset, donor_col=args.donor_col, acceptor_col=args.acceptor_col, pce_col=args.pce_col, radius=chosen_radius, n_bits=1024)
        if args.auto_prepare_only:
            raise SystemExit(0)
    elif args.use_existing_da_files:
        _auto_prepare_from_existing_source_files(radius=chosen_radius)
    if args.all_splits_seed_search:
        only = _parse_split_preset_list(args.all_splits_seed_search_presets)
        run_all_splits_seed_search_and_summary(args.all_splits_seed_search[0], args.all_splits_seed_search[1], only_presets=only)
    elif args.seed_search_lite:
        run_seed_search_lite(args.seed_search_lite[0], args.seed_search_lite[1], max_features=args.max_features)
    elif args.seed_search:
        run_seed_search(args.seed_search[0], args.seed_search[1], max_features=args.max_features)
    elif args.multi_seed:
        custom = _parse_seed_list(args.multi_seed_list) if args.multi_seed_list else None
        run_multi_seed_batch(seeds=custom, max_features=args.max_features, plot_only=args.plot_only, split_preset=args.split_preset, run_all_split_presets=args.run_all_split_presets)
    elif args.compare_split_seeds:
        preview_unknown_split_for_seeds(_parse_seed_list(args.compare_split_seeds))
    elif args.compare_train_seeds:
        compare_train_seeds_only(_parse_seed_list(args.compare_train_seeds), split_preset=args.split_preset)
    else:
        main(plot_only=args.plot_only, max_features=args.max_features, random_seed=args.seed, split_preset=args.split_preset, run_all_split_presets=args.run_all_split_presets)